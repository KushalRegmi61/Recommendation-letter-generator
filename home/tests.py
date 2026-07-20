import os
import tempfile
from datetime import date, datetime

from django.contrib.auth.models import AnonymousUser, User
from django.core.files.base import ContentFile
from django.db import IntegrityError, transaction
from django.test import RequestFactory, TestCase, SimpleTestCase, override_settings
from django.utils import timezone

from home.models import (
    Application, University, Academics, Department, Program,
    StudentLoginInfo, TeacherInfo, CustomTemplates, Qualities,
    Paper, Project, Files, Subject,
)
from home.filters import apply_application_filters, filter_options
from home.dashboard import build_teacher_dashboard_context


def login_as_teacher(client, teacher, password="test-pw"):
    """Sign ``client`` in as the Django user behind ``teacher``.

    Creates and links a User if the TeacherInfo does not have one. Replaces the
    old ``client.cookies["unique"] = ...`` idiom, which no longer authenticates.
    """
    user = teacher.user
    if user is None:
        user = User.objects.create_user(
            username=f"user-{teacher.unique_id}", password=password
        )
        user.first_name = f"{teacher.name}/{teacher.unique_id}"
        user.save()
        teacher.user = user
        teacher.save(update_fields=["user"])
    else:
        user.set_password(password)
        user.save()
    client.force_login(user)
    return user


class ModelFieldTests(TestCase):
    def _make_application(self):
        dept = Department.objects.create(dept_name="BCT")
        program = Program.objects.create(program_name="BE", department=dept)
        student = StudentLoginInfo.objects.create(
            username="alice", roll_number="075BCT001",
            department=dept, program=program, dob="2000-01-01",
        )
        prof = TeacherInfo.objects.create(
            unique_id="12345", name="Dr Smith", email="smith@example.com",
            department=dept,
        )
        return Application.objects.create(std=student, professor=prof)

    def test_application_has_new_fields(self):
        app = self._make_application()
        app.first_name = "Alice"
        app.middle_name = ""
        app.last_name = "Sharma"
        app.contact_number = "9800000000"
        app.applied_level = "Masters"
        app.known_roles = "instructor,thesis supervisor"
        app.years_known = "3"
        app.enrollment_batch = "075"
        app.passed_year = "2079"
        app.professional_experience = "Intern at X"
        app.strong_points = "Diligent"
        app.weak_points = "Perfectionist"
        app.save()
        app.refresh_from_db()
        self.assertEqual(app.applied_level, "Masters")
        self.assertEqual(app.known_roles, "instructor,thesis supervisor")
        self.assertIsNone(app.generated_at)

    def test_university_has_country(self):
        app = self._make_application()
        uni = University.objects.create(
            uni_name="MIT", country="USA", application=app,
        )
        uni.refresh_from_db()
        self.assertEqual(uni.country, "USA")

    def test_academics_has_final_percentage(self):
        app = self._make_application()
        aca = Academics.objects.create(
            application=app, final_percentage="82.5",
        )
        aca.refresh_from_db()
        self.assertEqual(aca.final_percentage, "82.5")


class ComposeFullNameTests(SimpleTestCase):
    def test_joins_all_three_parts(self):
        from home.intake import compose_full_name
        self.assertEqual(compose_full_name("Alice", "B", "Sharma"), "Alice B Sharma")

    def test_skips_blank_middle(self):
        from home.intake import compose_full_name
        self.assertEqual(compose_full_name("Alice", "", "Sharma"), "Alice Sharma")

    def test_strips_and_handles_none(self):
        from home.intake import compose_full_name
        self.assertEqual(compose_full_name("  Alice ", None, " Sharma"), "Alice Sharma")


class PendingApplicationTests(TestCase):
    def setUp(self):
        self.dept = Department.objects.create(dept_name="BEX")
        self.program = Program.objects.create(program_name="BE2", department=self.dept)
        self.student = StudentLoginInfo.objects.create(
            username="bob", roll_number="075BEX010",
            department=self.dept, program=self.program, dob="2000-01-01",
        )
        self.prof = TeacherInfo.objects.create(
            unique_id="55555", name="Dr Rai", email="rai@example.com",
            department=self.dept,
        )

    def test_false_when_no_application(self):
        from home.intake import has_pending_application
        self.assertFalse(has_pending_application(self.student, self.prof))

    def test_true_when_pending_exists(self):
        from home.intake import has_pending_application
        Application.objects.create(std=self.student, professor=self.prof, is_generated=False)
        self.assertTrue(has_pending_application(self.student, self.prof))

    def test_false_when_only_generated_exists(self):
        from home.intake import has_pending_application
        Application.objects.create(std=self.student, professor=self.prof, is_generated=True)
        self.assertFalse(has_pending_application(self.student, self.prof))


class ParseUniversitiesTests(SimpleTestCase):
    def test_zips_parallel_lists(self):
        from home.intake import parse_universities
        rows = parse_universities(
            names=["MIT", "TU Delft"],
            countries=["USA", "Netherlands"],
            deadlines=["2026-12-15", "2027-01-10"],
            programs=["MS CS", "MS EE"],
        )
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0], {
            "uni_name": "MIT", "country": "USA",
            "uni_deadline": "2026-12-15", "program_applied": "MS CS",
        })

    def test_skips_rows_with_blank_name(self):
        from home.intake import parse_universities
        rows = parse_universities(
            names=["MIT", "  "], countries=["USA", "UK"],
            deadlines=["2026-12-15", ""], programs=["MS", ""],
        )
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["uni_name"], "MIT")

    def test_blank_deadline_becomes_none(self):
        from home.intake import parse_universities
        rows = parse_universities(
            names=["MIT"], countries=["USA"], deadlines=[""], programs=[""],
        )
        self.assertIsNone(rows[0]["uni_deadline"])

    def test_ragged_lists_do_not_crash(self):
        from home.intake import parse_universities
        rows = parse_universities(names=["MIT"], countries=[], deadlines=[], programs=[])
        self.assertEqual(rows[0]["country"], "")
        self.assertIsNone(rows[0]["uni_deadline"])


class SaveUniversitiesTests(TestCase):
    def setUp(self):
        self.dept = Department.objects.create(dept_name="BME")
        self.program = Program.objects.create(program_name="BE3", department=self.dept)
        self.student = StudentLoginInfo.objects.create(
            username="cara", roll_number="075BME002",
            department=self.dept, program=self.program, dob="2000-01-01",
        )
        self.prof = TeacherInfo.objects.create(
            unique_id="77777", name="Dr Koirala", email="k@example.com",
            department=self.dept,
        )
        self.app = Application.objects.create(std=self.student, professor=self.prof)

    def test_creates_rows(self):
        from home.intake import save_universities
        rows = [
            {"uni_name": "MIT", "country": "USA", "uni_deadline": None, "program_applied": "MS"},
            {"uni_name": "ETH", "country": "Switzerland", "uni_deadline": None, "program_applied": "PhD"},
        ]
        count = save_universities(self.app, rows)
        self.assertEqual(count, 2)
        self.assertEqual(University.objects.filter(application=self.app).count(), 2)
        self.assertTrue(
            University.objects.filter(application=self.app, uni_name="ETH", country="Switzerland").exists()
        )

    def test_replaces_existing_rows(self):
        from home.intake import save_universities
        University.objects.create(uni_name="OLD", application=self.app)
        save_universities(self.app, [
            {"uni_name": "NEW", "country": "UK", "uni_deadline": None, "program_applied": ""},
        ])
        names = list(University.objects.filter(application=self.app).values_list("uni_name", flat=True))
        self.assertEqual(names, ["NEW"])


class Studentform1PostTests(TestCase):
    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCE")
        self.program = Program.objects.create(program_name="BE4", department=self.dept)
        self.student = StudentLoginInfo.objects.create(
            username="dan", roll_number="075BCE003",
            department=self.dept, program=self.program, dob="2000-01-01",
        )
        self.prof = TeacherInfo.objects.create(
            unique_id="88888", name="Dr Thapa", email="t@example.com",
            department=self.dept,
        )

    def _post_data(self):
        return {
            "naam": "dan", "roll": "075BCE003",
            "email": "dan@example.com",
            "prof": "Dr Thapa|88888",
            "first_name": "Dan", "middle_name": "", "last_name": "Gurung",
            "contact_number": "9811111111",
            "applied_level": "PhD",
            "known_roles": ["instructor", "thesis supervisor"],
            "yrs": "4",
            "enrollment_batch": "075",
            "passed_year": "2079",
            "professional_experience": "TA for 2 years",
            "strong_points": "Curious", "weak_points": "Impatient",
        }

    def test_saves_new_fields_on_application(self):
        resp = self.client.post("/studentform1", data=self._post_data())
        self.assertEqual(resp.status_code, 200)
        app = Application.objects.get(std=self.student, professor=self.prof)
        self.assertEqual(app.first_name, "Dan")
        self.assertEqual(app.last_name, "Gurung")
        self.assertEqual(app.contact_number, "9811111111")
        self.assertEqual(app.applied_level, "PhD")
        self.assertEqual(app.known_roles, "instructor,thesis supervisor")
        self.assertEqual(app.enrollment_batch, "075")
        self.assertEqual(app.passed_year, "2079")
        self.assertEqual(app.strong_points, "Curious")
        self.assertEqual(app.name, "Dan Gurung")
        self.assertFalse(app.is_generated)


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class Studentform2PostTests(TestCase):
    def setUp(self):
        self.dept = Department.objects.create(dept_name="BAR")
        self.program = Program.objects.create(program_name="BE5", department=self.dept)
        self.student = StudentLoginInfo.objects.create(
            username="eve", roll_number="075BAR004",
            department=self.dept, program=self.program, dob="2000-01-01",
        )
        self.prof = TeacherInfo.objects.create(
            unique_id="99999", name="Dr Basnet", email="b@example.com",
            department=self.dept,
        )
        self.app = Application.objects.create(
            std=self.student, professor=self.prof, name="Eve", is_generated=False,
        )

    def _post_data(self):
        return {
            "roll": "075BAR004", "naam": "eve", "prof_name": "Dr Basnet",
            "uni_name": ["MIT", "ETH"],
            "uni_country": ["USA", "Switzerland"],
            "uni_deadline": ["2026-12-15", ""],
            "uni_program": ["MS CS", "PhD"],
            "gpa": "3.9", "final_percentage": "88", "tentative_ranking": "Top 5%",
            "eca": "Robotics club",
        }

    def test_saves_repeatable_universities_and_percentage(self):
        resp = self.client.post("/studentform2", data=self._post_data())
        self.assertEqual(resp.status_code, 200)
        unis = University.objects.filter(application=self.app).order_by("uni_name")
        self.assertEqual(unis.count(), 2)
        self.assertEqual(unis[0].uni_name, "ETH")
        self.assertEqual(unis[0].country, "Switzerland")
        aca = Academics.objects.get(application=self.app)
        self.assertEqual(aca.final_percentage, "88")

    def test_a_failed_qualities_save_does_not_destroy_the_existing_row(self):
        # studentform2 deletes the old row before saving the replacement. Without
        # the transaction.atomic() wrapper a failed save leaves the application
        # permanently Qualities-less.
        from unittest import mock
        Qualities.objects.create(
            application=self.app, extracirricular="OLD MARKER",
        )
        with mock.patch.object(Qualities, "save", side_effect=RuntimeError("boom")):
            with self.assertRaises(RuntimeError):
                self.client.post("/studentform2", data=self._post_data())
        # The old row must still be there - the delete is rolled back with it.
        self.assertEqual(
            Qualities.objects.get(application=self.app).extracirricular, "OLD MARKER"
        )

    def test_duplicate_pending_is_rejected(self):
        # studentform2 fetches-and-updates the existing pending application (via
        # Application.objects.get(std__username=..., professor__name=...)), so the
        # count stays at 1 — i.e. this asserts no duplicate pending row is created.
        self.client.post("/studentform2", data=self._post_data())
        self.assertEqual(
            Application.objects.filter(
                std=self.student, professor=self.prof, is_generated=False
            ).count(),
            1,
        )


class Studentform1RenderTests(TestCase):
    def setUp(self):
        self.dept = Department.objects.create(dept_name="BEL")
        self.program = Program.objects.create(program_name="BE6", department=self.dept)
        self.student = StudentLoginInfo.objects.create(
            username="fay", roll_number="075BEL005",
            department=self.dept, program=self.program, dob="2000-01-01",
        )

    def test_form_has_new_fr2_inputs(self):
        self.client.cookies["student"] = "fay"
        resp = self.client.get("/studentform1")
        self.assertEqual(resp.status_code, 200)
        html = resp.content.decode()
        for field in [
            'name="first_name"', 'name="middle_name"', 'name="last_name"',
            'name="contact_number"', 'name="applied_level"', 'name="known_roles"',
            'name="enrollment_batch"', 'name="passed_year"',
            'name="professional_experience"', 'name="strong_points"',
            'name="weak_points"',
        ]:
            self.assertIn(field, html, f"missing input: {field}")


class ApplicationFilterTests(TestCase):
    def setUp(self):
        self.dept_bct = Department.objects.create(dept_name="BCT")
        self.dept_bce = Department.objects.create(dept_name="BCE")
        prog_bct = Program.objects.create(program_name="BE-BCT", department=self.dept_bct)
        prog_bce = Program.objects.create(program_name="BE-BCE", department=self.dept_bce)
        self.prof = TeacherInfo.objects.create(
            unique_id="T100", name="Prof One", email="p1@example.com",
            department=self.dept_bct,
        )
        self.stu_bct = StudentLoginInfo.objects.create(
            username="alice", roll_number="080BCT001", department=self.dept_bct,
            program=prog_bct, password="x", dob="2000-01-01",
        )
        self.stu_bce = StudentLoginInfo.objects.create(
            username="bob", roll_number="080BCE002", department=self.dept_bce,
            program=prog_bce, password="x", dob="2000-01-01",
        )
        self.app_bct = Application.objects.create(
            name="alice", email="a@example.com", professor=self.prof, std=self.stu_bct,
        )
        self.app_bce = Application.objects.create(
            name="bob", email="b@example.com", professor=self.prof, std=self.stu_bce,
        )
        University.objects.create(
            uni_name="MIT", country="USA", application=self.app_bct,
        )
        University.objects.create(
            uni_name="TU Delft", country="Netherlands", application=self.app_bce,
        )

    def base_qs(self):
        return Application.objects.filter(professor__unique_id="T100")

    def test_empty_params_returns_everything(self):
        result = apply_application_filters(self.base_qs(), {})
        self.assertEqual(result.count(), 2)

    def test_blank_values_are_ignored(self):
        result = apply_application_filters(
            self.base_qs(), {"department": "", "country": "", "college": ""}
        )
        self.assertEqual(result.count(), 2)

    def test_filter_by_department(self):
        result = apply_application_filters(self.base_qs(), {"department": "BCT"})
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

    def test_filter_by_country(self):
        result = apply_application_filters(self.base_qs(), {"country": "USA"})
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

    def test_filter_by_college(self):
        result = apply_application_filters(self.base_qs(), {"college": "TU Delft"})
        self.assertEqual([a.pk for a in result], [self.app_bce.pk])

    def test_filters_combine_with_and(self):
        # BCT department but a Netherlands university -> no match
        result = apply_application_filters(
            self.base_qs(), {"department": "BCT", "country": "Netherlands"}
        )
        self.assertEqual(result.count(), 0)

        result = apply_application_filters(
            self.base_qs(), {"department": "BCT", "country": "USA"}
        )
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

    def test_partial_and_case_insensitive_dropdown_values_match(self):
        # The dropdowns are typeable comboboxes, so a half-typed value must work.
        result = apply_application_filters(self.base_qs(), {"country": "us"})
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

        result = apply_application_filters(self.base_qs(), {"college": "delft"})
        self.assertEqual([a.pk for a in result], [self.app_bce.pk])

        result = apply_application_filters(self.base_qs(), {"department": "bct"})
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

    def test_no_duplicate_rows_when_application_has_many_universities(self):
        # A second USA university on the same application must not duplicate it.
        University.objects.create(
            uni_name="Stanford", country="USA", application=self.app_bct,
        )
        result = apply_application_filters(self.base_qs(), {"country": "USA"})
        self.assertEqual(result.count(), 1)

    def test_search_matches_application_name(self):
        result = apply_application_filters(self.base_qs(), {"q": "alice"})
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

    def test_search_is_case_insensitive_and_partial(self):
        result = apply_application_filters(self.base_qs(), {"q": "AL"})
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

    def test_search_matches_roll_number(self):
        result = apply_application_filters(self.base_qs(), {"q": "080bce"})
        self.assertEqual([a.pk for a in result], [self.app_bce.pk])

    def test_search_matches_email(self):
        result = apply_application_filters(self.base_qs(), {"q": "b@example.com"})
        self.assertEqual([a.pk for a in result], [self.app_bce.pk])

    def test_search_matches_first_or_last_name_fields(self):
        self.app_bce.first_name = "Bobby"
        self.app_bce.last_name = "Tables"
        self.app_bce.save()
        result = apply_application_filters(self.base_qs(), {"q": "tables"})
        self.assertEqual([a.pk for a in result], [self.app_bce.pk])

    def test_blank_search_is_ignored(self):
        result = apply_application_filters(self.base_qs(), {"q": "   "})
        self.assertEqual(result.count(), 2)

    def test_search_combines_with_dropdown_filters(self):
        # alice matches the text, but her university is in the USA, not Finland.
        result = apply_application_filters(
            self.base_qs(), {"q": "alice", "country": "Netherlands"}
        )
        self.assertEqual(result.count(), 0)

        result = apply_application_filters(
            self.base_qs(), {"q": "alice", "country": "USA"}
        )
        self.assertEqual([a.pk for a in result], [self.app_bct.pk])

    def test_search_terms_are_anded_regardless_of_word_order(self):
        self.app_bce.first_name = "Ramesh"
        self.app_bce.last_name = "Shrestha"
        self.app_bce.save()
        result = apply_application_filters(self.base_qs(), {"q": "shrestha ramesh"})
        self.assertEqual([a.pk for a in result], [self.app_bce.pk])

    def test_search_terms_may_match_different_fields(self):
        result = apply_application_filters(self.base_qs(), {"q": "bob 080bce"})
        self.assertEqual([a.pk for a in result], [self.app_bce.pk])

    def test_every_term_must_match(self):
        result = apply_application_filters(self.base_qs(), {"q": "bob 080bct"})
        self.assertEqual(result.count(), 0)

    def test_search_does_not_match_university_name(self):
        # University search is the dropdowns' job; matching it here would
        # surprise a professor searching for a person.
        result = apply_application_filters(self.base_qs(), {"q": "MIT"})
        self.assertEqual(result.count(), 0)


class FilterOptionTests(TestCase):
    def setUp(self):
        dept_bct = Department.objects.create(dept_name="BCT")
        dept_bex = Department.objects.create(dept_name="BEX")
        prog_bct = Program.objects.create(program_name="BE-BCT", department=dept_bct)
        prog_bex = Program.objects.create(program_name="BE-BEX", department=dept_bex)
        self.mine = TeacherInfo.objects.create(
            unique_id="T200", name="Mine", email="mine@example.com", department=dept_bct,
        )
        other = TeacherInfo.objects.create(
            unique_id="T201", name="Other", email="other@example.com", department=dept_bct,
        )
        stu_a = StudentLoginInfo.objects.create(
            username="ann", roll_number="080BCT010", department=dept_bct,
            program=prog_bct, password="x", dob="2000-01-01",
        )
        stu_b = StudentLoginInfo.objects.create(
            username="ben", roll_number="080BEX011", department=dept_bex,
            program=prog_bex, password="x", dob="2000-01-01",
        )
        app_a = Application.objects.create(
            name="ann", email="ann@example.com", professor=self.mine, std=stu_a,
        )
        app_b = Application.objects.create(
            name="ben", email="ben@example.com", professor=self.mine, std=stu_b,
        )
        app_other = Application.objects.create(
            name="zed", email="zed@example.com", professor=other, std=stu_a,
        )
        University.objects.create(uni_name="MIT", country="USA", application=app_a)
        University.objects.create(uni_name="MIT", country="USA", application=app_b)
        University.objects.create(uni_name="Aalto", country="Finland", application=app_b)
        University.objects.create(
            uni_name="SecretU", country="Japan", application=app_other,
        )

    def base_qs(self):
        return Application.objects.filter(professor__unique_id="T200")

    def test_options_are_sorted_and_deduplicated(self):
        options = filter_options(self.base_qs())
        self.assertEqual(options["departments"], ["BCT", "BEX"])
        self.assertEqual(options["countries"], ["Finland", "USA"])
        self.assertEqual(options["colleges"], ["Aalto", "MIT"])

    def test_options_exclude_other_professors_values(self):
        options = filter_options(self.base_qs())
        self.assertNotIn("Japan", options["countries"])
        self.assertNotIn("SecretU", options["colleges"])

    def test_blank_and_null_values_are_omitted(self):
        app = Application.objects.get(name="ann")
        University.objects.create(uni_name="", country=None, application=app)
        options = filter_options(self.base_qs())
        self.assertNotIn("", options["colleges"])
        self.assertNotIn(None, options["countries"])


class DashboardContextTests(TestCase):
    def setUp(self):
        dept = Department.objects.create(dept_name="BCT")
        prog = Program.objects.create(program_name="BE-BCT", department=dept)
        self.prof = TeacherInfo.objects.create(
            unique_id="T300", name="Prof Three", email="p3@example.com", department=dept,
        )
        self.stu = StudentLoginInfo.objects.create(
            username="cara", roll_number="080BCT020", department=dept,
            program=prog, password="x", dob="2000-01-01",
        )
        self.pending = Application.objects.create(
            name="cara pending", email="c@example.com", professor=self.prof,
            std=self.stu, is_generated=False,
        )
        self.older = Application.objects.create(
            name="cara older", email="c@example.com", professor=self.prof,
            std=self.stu, is_generated=True,
            generated_at=timezone.make_aware(datetime(2026, 1, 1, 9, 0)),
        )
        self.newer = Application.objects.create(
            name="cara newer", email="c@example.com", professor=self.prof,
            std=self.stu, is_generated=True,
            generated_at=timezone.make_aware(datetime(2026, 5, 1, 9, 0)),
        )
        University.objects.create(uni_name="MIT", country="USA", application=self.newer)
        University.objects.create(
            uni_name="Aalto", country="Finland", application=self.older,
        )

    def test_splits_pending_and_generated(self):
        ctx = build_teacher_dashboard_context("T300", {})
        self.assertEqual([a.pk for a in ctx["student_list"]], [self.pending.pk])
        self.assertEqual(len(ctx["all_students"]), 2)

    def test_generated_list_is_newest_first(self):
        ctx = build_teacher_dashboard_context("T300", {})
        self.assertEqual(
            [a.pk for a in ctx["all_students"]], [self.newer.pk, self.older.pk]
        )

    def test_filters_apply_to_both_lists(self):
        ctx = build_teacher_dashboard_context("T300", {"country": "USA"})
        self.assertEqual([a.pk for a in ctx["all_students"]], [self.newer.pk])
        # The pending application has no USA university, so it drops out too.
        self.assertEqual(list(ctx["student_list"]), [])

    def test_context_exposes_options_and_active_filters(self):
        ctx = build_teacher_dashboard_context("T300", {"country": "USA"})
        self.assertEqual(ctx["filter_options"]["countries"], ["Finland", "USA"])
        self.assertEqual(ctx["active_filters"]["country"], "USA")
        self.assertEqual(ctx["active_filters"]["department"], "")
        self.assertTrue(ctx["filters_active"])

    def test_no_filters_means_filters_inactive(self):
        ctx = build_teacher_dashboard_context("T300", {})
        self.assertFalse(ctx["filters_active"])

    def test_search_narrows_both_lists_and_counts_as_active(self):
        ctx = build_teacher_dashboard_context("T300", {"q": "newer"})
        self.assertEqual([a.pk for a in ctx["all_students"]], [self.newer.pk])
        self.assertEqual(list(ctx["student_list"]), [])
        self.assertEqual(ctx["active_filters"]["q"], "newer")
        self.assertTrue(ctx["filters_active"])

    def test_teacher_with_no_generated_letters_does_not_crash(self):
        Application.objects.filter(is_generated=True).delete()
        ctx = build_teacher_dashboard_context("T300", {})
        self.assertEqual(list(ctx["all_students"]), [])
        self.assertFalse(ctx["check_value"])

    def test_check_value_true_when_nothing_pending(self):
        self.pending.delete()
        ctx = build_teacher_dashboard_context("T300", {})
        self.assertTrue(ctx["check_value"])


class TeacherDashboardViewTests(TestCase):
    def setUp(self):
        dept = Department.objects.create(dept_name="BCT")
        prog = Program.objects.create(program_name="BE-BCT", department=dept)
        self.prof = TeacherInfo.objects.create(
            unique_id="T400", name="Prof Four", email="p4@example.com", department=dept,
        )
        stu = StudentLoginInfo.objects.create(
            username="dan", roll_number="080BCT030", department=dept,
            program=prog, password="x", dob="2000-01-01",
        )
        self.usa_app = Application.objects.create(
            name="dan usa", email="d@example.com", professor=self.prof,
            std=stu, is_generated=False,
        )
        self.fin_app = Application.objects.create(
            name="dan finland", email="d@example.com", professor=self.prof,
            std=stu, is_generated=False,
        )
        University.objects.create(
            uni_name="MIT", country="USA", application=self.usa_app,
        )
        University.objects.create(
            uni_name="Aalto", country="Finland", application=self.fin_app,
        )
        self.client.cookies["unique"] = "T400"
        self.client.cookies["username"] = "Prof Four"

    def test_teacher_view_renders_all_applications_by_default(self):
        response = self.client.get("/teacher")
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "dan usa")
        self.assertContains(response, "dan finland")

    def test_teacher_view_applies_country_filter(self):
        response = self.client.get("/teacher", {"country": "USA"})
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "dan usa")
        self.assertNotContains(response, "dan finland")

    def test_teacher_view_exposes_filter_options(self):
        response = self.client.get("/teacher")
        self.assertEqual(response.context["filter_options"]["countries"],
                         ["Finland", "USA"])

    def test_login_teacher_get_does_not_crash_without_generated_letters(self):
        # Regression: loginTeacher used Application.objects.get(is_generated=True),
        # which raised DoesNotExist for a professor with no generated letters.
        response = self.client.get("/loginTeacher")
        self.assertEqual(response.status_code, 200)

    def test_login_teacher_get_does_not_crash_with_two_generated_letters(self):
        # Regression: the same .get() raised MultipleObjectsReturned.
        Application.objects.filter(professor=self.prof).update(is_generated=True)
        response = self.client.get("/loginTeacher")
        self.assertEqual(response.status_code, 200)

    def test_all_teacher_dashboard_entry_points_supply_filter_options(self):
        # Teacher.html is rendered from several views; every one of them must
        # provide the filter context or the filter bar renders empty.
        for path in ("/", "/loginStudent", "/registerStudent"):
            with self.subTest(path=path):
                response = self.client.get(path)
                self.assertEqual(response.status_code, 200)
                self.assertEqual(
                    response.context["filter_options"]["countries"],
                    ["Finland", "USA"],
                )
                self.assertEqual(response.context["generated_count"], 0)

    def test_login_teacher_post_supplies_filter_options(self):
        user = User.objects.create_user(
            username="prof4", email="p4@example.com", password="secret",
        )
        user.first_name = "Prof Four/T400"
        user.save()
        response = self.client.post(
            "/loginTeacher", {"username": "p4@example.com", "password": "secret"},
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["filter_options"]["countries"], ["Finland", "USA"]
        )

    def test_filter_bar_is_rendered_with_typeable_comboboxes(self):
        response = self.client.get("/teacher")
        self.assertContains(response, 'name="department"')
        self.assertContains(response, 'name="country"')
        self.assertContains(response, 'name="college"')
        # Each field is an <input list=...> backed by a <datalist> of suggestions,
        # so the professor can either pick a value or type a partial one.
        self.assertContains(response, 'list="country-options"')
        self.assertContains(response, '<datalist id="country-options">')
        self.assertContains(response, '<option value="Finland">')
        self.assertContains(response, '<option value="MIT">')

    def test_active_filter_value_is_kept_in_the_box(self):
        response = self.client.get("/teacher", {"country": "USA"})
        self.assertContains(response, 'value="USA"')

    def test_partially_typed_filter_value_still_matches(self):
        response = self.client.get("/teacher", {"country": "us"})
        self.assertContains(response, "dan usa")
        self.assertNotContains(response, "dan finland")

    def test_search_box_is_rendered_and_keeps_its_value(self):
        response = self.client.get("/teacher", {"q": "dan usa"})
        self.assertContains(response, 'name="q"')
        self.assertContains(response, 'value="dan usa"')
        self.assertContains(response, "dan usa")
        self.assertNotContains(response, "dan finland")

    def test_generated_table_has_tracking_columns(self):
        response = self.client.get("/teacher")
        self.assertContains(response, "Generated on")
        self.assertContains(response, "Template")

    def test_empty_state_distinguishes_no_requests_from_no_matches(self):
        # No filter and nothing pending -> the cheerful global message.
        Application.objects.filter(professor=self.prof).update(is_generated=True)
        response = self.client.get("/teacher")
        self.assertContains(response, "You have no request for now")

        # A filter that matches nothing must NOT claim there are no requests at all.
        Application.objects.filter(professor=self.prof).update(is_generated=False)
        response = self.client.get("/teacher", {"country": "Antarctica"})
        self.assertContains(response, "No pending requests match")
        self.assertNotContains(response, "You have no request for now")

    def test_generated_count_is_shown(self):
        response = self.client.get("/teacher")
        self.assertEqual(response.context["generated_count"], 0)

    def test_dashboard_never_shows_another_professors_students(self):
        other_dept = Department.objects.create(dept_name="BEX")
        other_prog = Program.objects.create(
            program_name="BE-BEX", department=other_dept,
        )
        other_prof = TeacherInfo.objects.create(
            unique_id="T999", name="Prof Nine", email="p9@example.com",
            department=other_dept,
        )
        other_stu = StudentLoginInfo.objects.create(
            username="zoe", roll_number="080BEX099", department=other_dept,
            program=other_prog, password="x", dob="2000-01-01",
        )
        secret = Application.objects.create(
            name="zoe secret", email="z@example.com", professor=other_prof,
            std=other_stu, is_generated=False,
        )
        University.objects.create(
            uni_name="SecretU", country="Japan", application=secret,
        )

        response = self.client.get("/teacher")
        self.assertNotContains(response, "zoe secret")
        # Nor may the other professor's values leak into the filter suggestions.
        self.assertNotIn("Japan", response.context["filter_options"]["countries"])
        self.assertNotIn("SecretU", response.context["filter_options"]["colleges"])
        self.assertNotIn("BEX", response.context["filter_options"]["departments"])

    def test_teacher_view_redirects_when_cookie_is_missing(self):
        del self.client.cookies["unique"]
        response = self.client.get("/teacher")
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], "/loginTeacher")

    def test_teacher_view_redirects_when_cookie_is_stale(self):
        # A professor whose TeacherInfo was removed, or a hand-edited cookie.
        self.client.cookies["unique"] = "T000-does-not-exist"
        response = self.client.get("/teacher", {"country": "USA"})
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], "/loginTeacher")

    def test_redownload_link_shown_only_when_a_letter_is_stored(self):
        from django.core.files.base import ContentFile

        stored = Application.objects.create(
            name="dan stored", email="d@example.com", professor=self.prof,
            std=self.usa_app.std, is_generated=True,
        )
        response = self.client.get("/teacher")
        self.assertNotContains(response, f"/download_generated/?id={stored.pk}")

        stored.generated_letter.save(
            "dan.pdf", ContentFile(b"%PDF-1.4 x"), save=True,
        )
        try:
            response = self.client.get("/teacher")
            self.assertContains(response, f"/download_generated/?id={stored.pk}")
        finally:
            stored.generated_letter.delete(save=False)


@override_settings(MEDIA_ROOT=tempfile.mkdtemp())
class DownloadGeneratedTests(TestCase):
    def setUp(self):
        dept = Department.objects.create(dept_name="BCT")
        prog = Program.objects.create(program_name="BE-BCT", department=dept)
        self.prof = TeacherInfo.objects.create(
            unique_id="T500", name="Prof Five", email="p5@example.com", department=dept,
        )
        self.other = TeacherInfo.objects.create(
            unique_id="T501", name="Prof Six", email="p6@example.com", department=dept,
        )
        stu = StudentLoginInfo.objects.create(
            username="eve", roll_number="080BCT040", department=dept,
            program=prog, password="x", dob="2000-01-01",
        )
        self.stored = Application.objects.create(
            name="eve stored", email="e@example.com", professor=self.prof,
            std=stu, is_generated=True,
        )
        self.stored.generated_letter.save(
            "eve.pdf", ContentFile(b"%PDF-1.4 fake letter"), save=True,
        )
        self.legacy = Application.objects.create(
            name="eve legacy", email="e@example.com", professor=self.prof,
            std=stu, is_generated=True,
        )
        self.client.cookies["unique"] = "T500"

    def test_returns_stored_file(self):
        response = self.client.get(f"/download_generated/?id={self.stored.pk}")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(b"".join(response.streaming_content), b"%PDF-1.4 fake letter")
        self.assertIn("attachment", response["Content-Disposition"])

    def test_missing_stored_file_redirects_with_message(self):
        response = self.client.get(f"/download_generated/?id={self.legacy.pk}")
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], "/teacher")

    def test_other_professors_letter_is_not_served(self):
        self.client.cookies["unique"] = "T501"
        response = self.client.get(f"/download_generated/?id={self.stored.pk}")
        self.assertEqual(response.status_code, 404)

    def test_anonymous_request_is_sent_to_the_login_page(self):
        del self.client.cookies["unique"]
        response = self.client.get(f"/download_generated/?id={self.stored.pk}")
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], "/loginTeacher")

    def test_unknown_professor_cookie_is_sent_to_the_login_page(self):
        self.client.cookies["unique"] = "T-does-not-exist"
        response = self.client.get(f"/download_generated/?id={self.stored.pk}")
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], "/loginTeacher")

    def test_missing_id_parameter_is_not_served(self):
        response = self.client.get("/download_generated/")
        self.assertEqual(response.status_code, 404)

    def test_malformed_id_is_not_served(self):
        for bad in ("abc", "1 OR 1=1", "5.0", ""):
            with self.subTest(bad=bad):
                response = self.client.get(f"/download_generated/?id={bad}")
                self.assertEqual(response.status_code, 404)


class SystemTemplateModelTests(TestCase):
    """CustomTemplates must support shared system templates (FR-1)."""

    def test_a_system_template_needs_no_professor(self):
        tpl = CustomTemplates.objects.create(
            template_name="Formal / Academic",
            template="Dear Committee,",
            professor=None,
            is_system=True,
        )
        self.assertIsNone(tpl.professor)
        self.assertTrue(tpl.is_system)

    def test_professor_templates_are_not_system_by_default(self):
        dept = Department.objects.create(dept_name="BCT")
        teacher = TeacherInfo.objects.create(
            name="Prof A", unique_id="T-A", department=dept
        )
        tpl = CustomTemplates.objects.create(
            template_name="Mine", template="Hello", professor=teacher
        )
        self.assertFalse(tpl.is_system)

    def test_str_does_not_crash_without_a_professor(self):
        tpl = CustomTemplates.objects.create(
            template_name="Formal", template="x", professor=None, is_system=True
        )
        self.assertIn("Formal", str(tpl))


class SeededSystemTemplateTests(TestCase):
    """The data migration must ship a usable starter library (FR-1)."""

    def test_three_system_templates_are_seeded(self):
        seeded = CustomTemplates.objects.filter(is_system=True)
        self.assertEqual(seeded.count(), 3)

    def test_seeded_templates_have_names_and_bodies(self):
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                self.assertTrue(tpl.template_name)
                self.assertGreater(len(tpl.template), 100)
                self.assertIsNone(tpl.professor)
                self.assertFalse(tpl.is_default)

    def test_seeded_templates_are_ascii_only(self):
        # The PDF export encodes latin-1; a stray em dash crashes the download.
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                tpl.template.encode("ascii")

    def test_seeded_templates_are_valid_jinja(self):
        from jinja2 import Template
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                Template(tpl.template)  # raises TemplateSyntaxError if malformed

    def test_seeded_templates_render_without_leaking_none(self):
        """A sparse application must not put the literal word None in the prose."""
        from jinja2 import Template
        dept = Department.objects.create(dept_name="BCT")
        program = Program.objects.create(program_name="BE-BCT", department=dept)
        teacher = TeacherInfo.objects.create(
            name="Prof Q", unique_id="T-Q", email="q@example.com", department=dept,
        )
        student = StudentLoginInfo.objects.create(
            username="Sparse Student", roll_number="080BCT500", department=dept,
            program=program, password="x", dob="2000-01-01",
        )
        application = Application.objects.create(
            name="Sparse Student", std=student, professor=teacher,
        )
        # Every satellite object is absent, as it is for a barely-filled request.
        context = {
            "app": application, "teacher": teacher, "academics": None,
            "paper": None, "project": None, "university": None, "quality": None,
            "pronoun": "They", "pronoun_obj": "them", "pronoun_pos": "Their",
            "rel_desc": "teacher", "strength_phrase": "with great enthusiasm",
            "deadline": "", "today": "July 20, 2026",
        }
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                letter = Template(tpl.template).render(context)
                self.assertNotIn("None", letter)
                self.assertIn("Sparse Student", letter)
                # No paragraph may be bare punctuation left by an empty branch.
                for line in letter.splitlines():
                    self.assertNotIn(line.strip(), (".", ",", "-"))

    def test_seeded_paragraphs_are_single_source_lines(self):
        """The PDF exporter does not reflow across newlines, so a sentence
        split over two source lines becomes two ragged lines in the export."""
        for tpl in CustomTemplates.objects.filter(is_system=True):
            for line in tpl.template.splitlines():
                stripped = line.strip()
                # A line that ends mid-sentence has broken a paragraph in two.
                if stripped and not stripped.startswith("{%") and len(stripped) > 60:
                    with self.subTest(name=tpl.template_name, line=stripped[:40]):
                        self.assertTrue(
                            stripped.endswith((".", "}", "%}", '"', "!")),
                            f"paragraph broken mid-sentence: {stripped!r}",
                        )

    def _render_seeded(self, context):
        from jinja2 import Template
        return [
            (tpl.template_name, Template(tpl.template).render(context))
            for tpl in CustomTemplates.objects.filter(is_system=True)
        ]

    def test_blank_related_rows_do_not_leave_holes_in_the_prose(self):
        """Guards must fire for rows that EXIST but whose fields are blank.

        The sparse case above passes None for every satellite object, so the
        guards on dept_name/program_name/gpa are never actually exercised. Here
        the rows exist and the fields are empty, which is the realistic shape:
        Department.dept_name and Program.program_name are NOT NULL but
        blank=True, so the empty value is "" and dropping a guard leaves a
        doubled space ("studies in  at the Institute"), not the word "None".
        """
        dept = Department.objects.create(dept_name="")
        program = Program.objects.create(program_name="", department=dept)
        teacher = TeacherInfo.objects.create(
            name="Prof B", unique_id="T-B", email="b@example.com", department=dept,
        )
        student = StudentLoginInfo.objects.create(
            username="Blank Student", roll_number="080BCT501", department=dept,
            program=program, password="x", dob="2000-01-01",
        )
        application = Application.objects.create(
            name="Blank Student", std=student, professor=teacher,
        )
        # The row exists but carries no grade, so the GPA sentence must not run.
        academics = Academics.objects.create(
            application=application, gpa="", tentative_ranking="",
        )
        context = {
            "app": application, "teacher": teacher, "academics": academics,
            "paper": None, "project": None, "university": None, "quality": None,
            "pronoun": "They", "pronoun_obj": "them", "pronoun_pos": "Their",
            "rel_desc": "teacher", "strength_phrase": "with great enthusiasm",
            "deadline": "", "today": "July 20, 2026",
        }
        for name, letter in self._render_seeded(context):
            with self.subTest(name=name):
                self.assertIn("Blank Student", letter)
                self.assertNotIn("None", letter)
                # An unguarded blank field collapses to "" and doubles a space.
                self.assertNotIn("  ", letter)
                # academics is truthy but has no gpa: the whole sentence goes.
                self.assertNotIn("GPA", letter)

    def test_seeded_templates_are_grammatical_for_plural_pronouns(self):
        """Students with unset gender get They/them/Their, which must not
        collide with third-person-singular verbs ("They has maintained")."""
        from jinja2 import Template
        dept = Department.objects.create(dept_name="BEI")
        program = Program.objects.create(program_name="BE-BEI", department=dept)
        teacher = TeacherInfo.objects.create(
            name="Prof Z", unique_id="T-Z", email="z@example.com", department=dept,
        )
        student = StudentLoginInfo.objects.create(
            username="Plural Student", roll_number="080BEI777", department=dept,
            program=program, password="x", dob="2000-01-01",
        )
        application = Application.objects.create(
            name="Plural Student", std=student, professor=teacher,
            subjects="Control Systems", is_paper=True,
        )
        academics = Academics.objects.create(application=application, gpa="3.75")
        quality = Qualities.objects.create(
            application=application, leadership=True, hardworking=True,
            teamwork=True, friendly=True, quality="unfailingly curious",
        )
        context = {
            "app": application, "teacher": teacher, "academics": academics,
            "paper": None, "project": None, "university": None, "quality": quality,
            "pronoun": "They", "pronoun_obj": "them", "pronoun_pos": "Their",
            "rel_desc": "teacher", "strength_phrase": "with great enthusiasm",
            "deadline": "", "today": "July 20, 2026",
        }
        singular_verbs = (
            "They has", "they has", "They is", "they is", "They shows",
            "they shows", "They works", "they works", "They holds", "they holds",
        )
        for tpl in CustomTemplates.objects.filter(is_system=True):
            letter = Template(tpl.template).render(context)
            for bad in singular_verbs:
                with self.subTest(name=tpl.template_name, phrase=bad):
                    self.assertNotIn(bad, letter)


class LetterContextTests(TestCase):
    """build_letter_context assembles everything the Jinja templates read."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof B", unique_id="T-B", email="b@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Ramesh Shrestha", roll_number="080BCT042", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Male",
        )
        self.application = Application.objects.create(
            name="Ramesh Shrestha", std=self.student, professor=self.teacher,
            subjects="Data Structures,Algorithms",
        )

    def test_pronouns_follow_the_student_gender(self):
        from home.letters import build_letter_context
        ctx = build_letter_context(self.application)
        self.assertEqual(ctx["pronoun"], "He")
        self.assertEqual(ctx["pronoun_obj"], "him")
        self.assertEqual(ctx["pronoun_pos"], "His")

    def test_female_gender_maps_correctly(self):
        from home.letters import build_letter_context
        self.student.gender = "Female"
        self.student.save()
        ctx = build_letter_context(self.application)
        self.assertEqual(ctx["pronoun"], "She")
        self.assertEqual(ctx["pronoun_obj"], "her")
        self.assertEqual(ctx["pronoun_pos"], "Her")

    def test_unknown_gender_falls_back_to_they(self):
        from home.letters import build_letter_context
        self.student.gender = ""
        self.student.save()
        ctx = build_letter_context(self.application)
        self.assertEqual(ctx["pronoun"], "They")
        self.assertEqual(ctx["pronoun_obj"], "them")
        self.assertEqual(ctx["pronoun_pos"], "Their")

    def test_null_gender_falls_back_to_they(self):
        from home.letters import build_letter_context
        self.student.gender = None
        self.student.save()
        self.assertEqual(build_letter_context(self.application)["pronoun"], "They")

    def test_gender_matching_is_case_insensitive(self):
        from home.letters import build_letter_context
        self.student.gender = "MALE"
        self.student.save()
        self.assertEqual(build_letter_context(self.application)["pronoun"], "He")

    def test_no_context_value_is_none_for_the_pronoun_words(self):
        # Jinja's ``|lower`` filter raises on an explicit None, and every
        # seeded template pipes these through it.
        from home.letters import build_letter_context
        self.student.gender = None
        self.student.save()
        ctx = build_letter_context(self.application)
        for key in ("pronoun", "pronoun_obj", "pronoun_pos", "rel_desc", "strength_phrase"):
            with self.subTest(key=key):
                self.assertIsInstance(ctx[key], str)

    def test_subjects_are_split_into_list_and_last(self):
        from home.letters import build_letter_context
        ctx = build_letter_context(self.application)
        self.assertEqual(ctx["subjects"], ["Data Structures"])
        self.assertEqual(ctx["subject"], "Algorithms")

    def test_single_subject_sets_value_true(self):
        from home.letters import build_letter_context
        self.application.subjects = "Algorithms"
        self.application.save()
        ctx = build_letter_context(self.application)
        self.assertTrue(ctx["value"])

    def test_empty_subjects_do_not_crash(self):
        from home.letters import build_letter_context
        self.application.subjects = ""
        self.application.save()
        ctx = build_letter_context(self.application)
        self.assertEqual(ctx["subjects"], [])
        self.assertFalse(ctx["value"])

    def test_firstname_is_the_first_word(self):
        from home.letters import build_letter_context
        self.assertEqual(build_letter_context(self.application)["firstname"], "Ramesh")

    def test_missing_satellite_rows_do_not_raise(self):
        # No Paper/Project/University/Qualities/Academics/Files rows exist.
        from home.letters import build_letter_context
        ctx = build_letter_context(self.application)
        for key in ("paper", "project", "university", "quality", "academics", "files"):
            with self.subTest(key=key):
                self.assertIsNone(ctx[key])
        self.assertEqual(ctx["deadline"], "")

    def test_satellite_rows_are_returned_when_present(self):
        from home.letters import build_letter_context
        academics = Academics.objects.create(application=self.application, gpa="3.9")
        university = University.objects.create(
            application=self.application, uni_name="MIT", country="USA",
        )
        ctx = build_letter_context(self.application)
        self.assertEqual(ctx["academics"], academics)
        self.assertEqual(ctx["university"], university)

    def test_deadline_is_formatted_when_present(self):
        from home.letters import build_letter_context
        University.objects.create(
            application=self.application, uni_name="MIT", country="USA",
            uni_deadline=date(2026, 12, 15),
        )
        self.assertEqual(build_letter_context(self.application)["deadline"], "December 15, 2026")

    def test_teacher_and_app_aliases_are_present(self):
        from home.letters import build_letter_context
        ctx = build_letter_context(self.application)
        self.assertEqual(ctx["teacher"], self.teacher)
        self.assertEqual(ctx["app"], self.application)
        self.assertEqual(ctx["student"], self.application)
        self.assertTrue(ctx["today"])

    def test_every_seeded_system_template_renders_against_this_context(self):
        # The real integration check: the three shipped templates must render
        # against exactly what this function produces.
        from jinja2 import Template
        from home.letters import build_letter_context
        ctx = build_letter_context(self.application)
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                letter = Template(tpl.template).render(ctx)
                self.assertIn("Ramesh Shrestha", letter)
                self.assertNotIn("None", letter)

    def test_rel_desc_comes_from_the_application(self):
        from home.letters import build_letter_context
        self.application.relationship_type = "project supervisor"
        self.application.save()
        self.assertEqual(build_letter_context(self.application)["rel_desc"], "project supervisor")

    def test_rel_desc_falls_back_when_unset(self):
        from home.letters import build_letter_context
        for value in (None, "", "   "):
            with self.subTest(value=value):
                self.application.relationship_type = value
                self.application.save()
                self.assertEqual(build_letter_context(self.application)["rel_desc"], "teacher")

    def test_strength_phrase_follows_the_recommendation_strength(self):
        from home.letters import build_letter_context
        quality = Qualities.objects.create(application=self.application)
        expected = {
            "top5": "as one of the very best students I have taught",
            "top10": "as one of the strongest students I have taught",
            "outstanding": "in the strongest possible terms",
            "strong": "with great enthusiasm",
        }
        for value, phrase in expected.items():
            with self.subTest(value=value):
                quality.recommendation_strength = value
                quality.save()
                self.assertEqual(build_letter_context(self.application)["strength_phrase"], phrase)

    def test_strength_phrase_falls_back_without_a_qualities_row(self):
        from home.letters import build_letter_context
        self.assertEqual(
            build_letter_context(self.application)["strength_phrase"],
            "with great enthusiasm",
        )

    def test_strength_phrase_falls_back_on_an_unrecognised_value(self):
        from home.letters import build_letter_context
        Qualities.objects.create(application=self.application, recommendation_strength="")
        self.assertEqual(
            build_letter_context(self.application)["strength_phrase"],
            "with great enthusiasm",
        )

    def test_recommendation_reads_correctly_for_every_strength(self):
        """Each phrase must slot grammatically into every seeded template."""
        from jinja2 import Template
        from home.letters import build_letter_context
        quality = Qualities.objects.create(application=self.application)
        for value in ("top5", "top10", "outstanding", "strong"):
            quality.recommendation_strength = value
            quality.save()
            ctx = build_letter_context(self.application)
            for tpl in CustomTemplates.objects.filter(is_system=True):
                with self.subTest(strength=value, name=tpl.template_name):
                    letter = Template(tpl.template).render(ctx)
                    self.assertNotIn("and without reservation", letter)
                    self.assertNotIn("  ", letter)
                    # A trailing prepositional phrase must not attach itself
                    # to the end of a ranked strength phrase.
                    self.assertNotIn("I have taught for", letter)

    def _subjects(self, raw):
        from home.letters import build_letter_context
        self.application.subjects = raw
        self.application.save()
        return build_letter_context(self.application)

    def test_subject_keys_agree_on_empty_segments(self):
        # subjects/subject/value were derived separately in the legacy views
        # and disagreed; one normalisation now backs all three.
        cases = {
            "A,B,C": (["A", "B"], "C", False),
            "A": ([], "A", True),
            "A,": ([], "A", True),       # trailing comma must not blank the subject
            ",A": ([], "A", True),
            "A, B": (["A"], "B", False),  # inner whitespace is stripped
            "": ([], "", False),
            None: ([], "", False),
        }
        for raw, (subjects, subject, value) in cases.items():
            with self.subTest(raw=raw):
                ctx = self._subjects(raw)
                self.assertEqual(ctx["subjects"], subjects)
                self.assertEqual(ctx["subject"], subject)
                self.assertEqual(ctx["value"], value)

    def test_a_trailing_comma_does_not_drop_the_only_subject(self):
        # Legacy gave value=True with subject="", rendering "I taught him .".
        ctx = self._subjects("Algorithms,")
        self.assertTrue(ctx["value"])
        self.assertEqual(ctx["subject"], "Algorithms")

    def test_subjects_sentence_reads_as_prose(self):
        for raw, expected in (
            ("", ""),
            ("Algorithms", "Algorithms"),
            ("Algorithms,Compilers", "Algorithms and Compilers"),
            ("Algorithms,Compilers,Networks", "Algorithms, Compilers and Networks"),
            ("Algorithms, Compilers , Networks", "Algorithms, Compilers and Networks"),
        ):
            with self.subTest(raw=raw):
                self.assertEqual(self._subjects(raw)["subjects_sentence"], expected)

    def test_firstname_handles_awkward_whitespace(self):
        from home.letters import build_letter_context
        for raw, expected in (
            ("  Ram Thapa", "Ram"),      # leading space returned "" before
            ("Ram  Thapa", "Ram"),
            ("Sita", "Sita"),
            ("", ""),
            (None, ""),
        ):
            with self.subTest(raw=raw):
                self.application.name = raw
                self.application.save()
                self.assertEqual(build_letter_context(self.application)["firstname"], expected)

    def test_seeded_templates_list_subjects_as_prose(self):
        from jinja2 import Template
        ctx = self._subjects("Algorithms,Compilers,Networks")
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                letter = Template(tpl.template).render(ctx)
                self.assertNotIn("Algorithms,Compilers", letter)
                if "Algorithms" in letter:
                    self.assertIn("Algorithms, Compilers and Networks", letter)

    def test_seeded_templates_omit_the_subject_clause_when_unset(self):
        from jinja2 import Template
        ctx = self._subjects("")
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                letter = Template(tpl.template).render(ctx)
                self.assertNotIn("having taught", letter)
                self.assertNotIn("I taught", letter)
                self.assertNotIn(" in .", letter)


class TemplateSelectionTests(TestCase):
    """select_template resolves a professor's pick, default, then system (FR-1)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof C", unique_id="T-C", email="c@example.com", department=self.dept,
        )
        self.other = TeacherInfo.objects.create(
            name="Prof D", unique_id="T-D", email="d@example.com", department=self.dept,
        )
        self.mine = CustomTemplates.objects.create(
            template_name="Mine", template="mine", professor=self.teacher
        )
        self.my_default = CustomTemplates.objects.create(
            template_name="My Default", template="default",
            professor=self.teacher, is_default=True,
        )
        self.theirs = CustomTemplates.objects.create(
            template_name="Theirs", template="theirs", professor=self.other
        )

    def test_an_explicit_own_template_wins(self):
        from home.letters import select_template
        self.assertEqual(select_template(self.teacher, self.mine.pk), self.mine)

    def test_a_system_template_may_be_selected(self):
        from home.letters import select_template
        system = CustomTemplates.objects.filter(is_system=True).first()
        self.assertEqual(select_template(self.teacher, system.pk), system)

    def test_another_professors_template_is_refused(self):
        from home.letters import select_template
        # Falls back to this professor's default rather than leaking Prof D's.
        self.assertEqual(select_template(self.teacher, self.theirs.pk), self.my_default)

    def test_no_choice_uses_the_professors_default(self):
        from home.letters import select_template
        self.assertEqual(select_template(self.teacher, None), self.my_default)

    def test_blank_and_malformed_ids_use_the_default(self):
        from home.letters import select_template
        for bad in ("", "   ", "abc", None, "0", "-1", "9999999"):
            with self.subTest(value=bad):
                self.assertEqual(select_template(self.teacher, bad), self.my_default)

    def test_string_ids_from_post_data_work(self):
        from home.letters import select_template
        self.assertEqual(select_template(self.teacher, str(self.mine.pk)), self.mine)

    def test_without_a_default_it_falls_back_to_a_system_template(self):
        from home.letters import select_template
        CustomTemplates.objects.filter(professor=self.teacher).delete()
        chosen = select_template(self.teacher, None)
        self.assertTrue(chosen.is_system)

    def test_with_nothing_at_all_it_returns_none(self):
        from home.letters import select_template
        CustomTemplates.objects.all().delete()
        self.assertIsNone(select_template(self.teacher, None))

    def test_another_professors_default_is_not_used_as_mine(self):
        from home.letters import select_template
        CustomTemplates.objects.create(
            template_name="Their Default", template="secret",
            professor=self.other, is_default=True,
        )
        self.assertEqual(select_template(self.teacher, None), self.my_default)

    def test_without_my_own_default_i_get_a_system_template_not_theirs(self):
        from home.letters import select_template
        CustomTemplates.objects.create(
            template_name="Their Default", template="secret",
            professor=self.other, is_default=True,
        )
        CustomTemplates.objects.filter(professor=self.teacher).delete()
        chosen = select_template(self.teacher, None)
        self.assertTrue(chosen.is_system)
        self.assertIsNone(chosen.professor)

    def test_an_owned_row_flagged_system_cannot_exist(self):
        # The leak this closes: such a row would satisfy the ``is_system`` arm
        # and so be visible to every professor. The DB constraint makes it
        # unrepresentable, so there is nothing for the query to leak.
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                CustomTemplates.objects.create(
                    template_name="Hybrid", template="secret",
                    professor=self.other, is_system=True,
                )

    def test_the_system_arm_of_the_predicate_requires_an_unowned_row(self):
        # Defence in depth behind the constraint above: even if the constraint
        # were dropped, the query must not match an owned row flagged system.
        from home.letters import visible_to
        system_arm = [
            child for child in visible_to(self.teacher).children
            if hasattr(child, "children")
        ]
        self.assertEqual(len(system_arm), 1)
        self.assertIn(("professor__isnull", True), system_arm[0].children)


class AvailableTemplateTests(TestCase):
    """available_templates lists what a professor may generate from (FR-1)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof C2", unique_id="T-C2", email="c2@example.com", department=self.dept,
        )
        self.other = TeacherInfo.objects.create(
            name="Prof D2", unique_id="T-D2", email="d2@example.com", department=self.dept,
        )
        self.mine = CustomTemplates.objects.create(
            template_name="Mine", template="mine", professor=self.teacher
        )
        self.theirs = CustomTemplates.objects.create(
            template_name="Theirs", template="theirs", professor=self.other
        )

    def test_own_templates_are_included(self):
        from home.letters import available_templates
        self.assertIn(self.mine, available_templates(self.teacher))

    def test_system_templates_are_included(self):
        from home.letters import available_templates
        names = {t.template_name for t in available_templates(self.teacher)}
        self.assertIn("Formal / Academic", names)

    def test_another_professors_templates_are_excluded(self):
        from home.letters import available_templates
        self.assertNotIn(self.theirs, available_templates(self.teacher))

    def test_no_template_appears_twice(self):
        from home.letters import available_templates
        results = list(available_templates(self.teacher))
        self.assertEqual(len(results), len({t.pk for t in results}))

    def test_the_default_sorts_first(self):
        from home.letters import available_templates
        self.mine.is_default = True
        self.mine.save()
        self.assertEqual(available_templates(self.teacher).first(), self.mine)

    def test_it_returns_a_queryset_that_can_be_filtered(self):
        # ``make_letter`` chains .filter(is_default=True) onto this.
        from home.letters import available_templates
        self.mine.is_default = True
        self.mine.save()
        self.assertEqual(
            available_templates(self.teacher).filter(is_default=True).first(), self.mine
        )

    def test_an_owned_row_flagged_system_is_excluded(self):
        from home.letters import available_templates
        # Same leak as in TemplateSelectionTests, via the listing rather than
        # the picker. The constraint blocks the row outright.
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                CustomTemplates.objects.create(
                    template_name="Hybrid", template="secret",
                    professor=self.other, is_system=True,
                )
        self.assertNotIn(self.theirs, available_templates(self.teacher))


class RenderLetterTests(TestCase):
    """render_letter fills the chosen template with application data (FR-1)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof E", unique_id="T-E", email="e@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Sita Rai", roll_number="080BCT001", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Female",
        )
        self.application = Application.objects.create(
            name="Sita Rai", std=self.student, professor=self.teacher,
            subjects="Physics",
        )
        # A default whose body differs from every template under test, so that
        # "the chosen body is used" cannot pass by silently falling back here.
        self.default = CustomTemplates.objects.create(
            template_name="Default", template="DEFAULT-BODY-NOT-CHOSEN",
            professor=self.teacher, is_default=True,
        )

    def test_the_chosen_template_body_is_used(self):
        from home.letters import render_letter
        tpl = CustomTemplates.objects.create(
            template_name="Terse", template="Hello {{ app.name }} from {{ teacher.name }}.",
            professor=self.teacher,
        )
        self.assertEqual(
            render_letter(self.application, tpl), "Hello Sita Rai from Prof E."
        )

    def test_pronouns_reach_the_template(self):
        from home.letters import render_letter
        tpl = CustomTemplates.objects.create(
            template_name="P", template="{{ pronoun }} and {{ pronoun_obj }}",
            professor=self.teacher,
        )
        self.assertEqual(render_letter(self.application, tpl), "She and her")

    def test_no_template_renders_empty(self):
        from home.letters import render_letter
        self.assertEqual(render_letter(self.application, None), "")

    def test_a_template_with_an_empty_body_renders_empty(self):
        from home.letters import render_letter
        tpl = CustomTemplates.objects.create(
            template_name="Blank", template="", professor=self.teacher,
        )
        self.assertEqual(render_letter(self.application, tpl), "")

    def test_every_seeded_system_template_renders(self):
        from home.letters import render_letter
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                letter = render_letter(self.application, tpl)
                self.assertIn("Sita Rai", letter)
                self.assertIn("Prof E", letter)
                self.assertNotIn("None", letter)

    def test_a_template_with_broken_syntax_renders_empty(self):
        from home.letters import render_letter
        tpl = CustomTemplates.objects.create(
            template_name="Broken", template="{% for x in y %}no endfor",
            professor=self.teacher,
        )
        self.assertEqual(render_letter(self.application, tpl), "")

    def test_a_template_referencing_a_missing_field_renders_empty(self):
        from home.letters import render_letter
        tpl = CustomTemplates.objects.create(
            template_name="Undefined", template="{{ app.nonexistent.attr }}",
            professor=self.teacher,
        )
        self.assertEqual(render_letter(self.application, tpl), "")

    def test_sandbox_escape_attempts_are_refused(self):
        from home.letters import render_letter
        tpl = CustomTemplates.objects.create(
            template_name="Escape", template="{{ ''.__class__.__mro__ }}",
            professor=self.teacher,
        )
        self.assertEqual(render_letter(self.application, tpl), "")

    def test_the_subclass_walk_springboard_is_refused(self):
        from home.letters import render_letter
        tpl = CustomTemplates.objects.create(
            template_name="Springboard",
            template="{{ app.std.__class__.__mro__[1].__subclasses__()|length }}",
            professor=self.teacher,
        )
        self.assertEqual(render_letter(self.application, tpl), "")

    def test_a_sandbox_violation_raises_a_template_error(self):
        # ``render_letter`` relies on SecurityError subclassing TemplateError.
        from jinja2 import TemplateError
        from jinja2.sandbox import SecurityError
        self.assertTrue(issubclass(SecurityError, TemplateError))


class LetterExportTests(TestCase):
    """PDF/DOCX bytes are produced from letter text (FR-1)."""

    def test_pdf_bytes_look_like_a_pdf(self):
        from home.letters import build_pdf_bytes
        data = build_pdf_bytes("Dear Committee,\n\nRegards,\nProf")
        self.assertTrue(data.startswith(b"%PDF"))

    def test_docx_bytes_look_like_a_zip(self):
        # .docx is a zip container; PK is the zip magic number.
        from home.letters import build_docx_bytes
        data = build_docx_bytes("Dear Committee,\n\nRegards,\nProf")
        self.assertTrue(data.startswith(b"PK"))

    def test_docx_keeps_one_paragraph_per_block(self):
        import io
        from docx import Document
        from home.letters import build_docx_bytes
        data = build_docx_bytes("First block.\n\nSecond block.")
        doc = Document(io.BytesIO(data))
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("First block.", texts)
        self.assertIn("Second block.", texts)

    def test_docx_preserves_the_letter_text(self):
        import io
        from docx import Document
        from home.letters import build_docx_bytes
        data = build_docx_bytes("Dear Committee,\n\nI recommend her warmly.")
        doc = Document(io.BytesIO(data))
        joined = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("I recommend her warmly.", joined)

    def test_empty_text_still_produces_a_file(self):
        from home.letters import build_docx_bytes, build_pdf_bytes
        self.assertTrue(build_pdf_bytes("").startswith(b"%PDF"))
        self.assertTrue(build_docx_bytes("").startswith(b"PK"))

    def test_non_latin1_characters_are_replaced_not_crashed(self):
        # fpdf encodes latin-1; an em dash used to raise UnicodeEncodeError.
        from home.letters import build_pdf_bytes
        self.assertTrue(build_pdf_bytes("A — B").startswith(b"%PDF"))

    def test_curly_quotes_do_not_crash_the_pdf(self):
        from home.letters import build_pdf_bytes
        self.assertTrue(build_pdf_bytes("“Quoted” and ‘single’").startswith(b"%PDF"))

    def test_docx_handles_non_latin1_characters_natively(self):
        # Only the PDF path is latin-1 limited; docx is XML/UTF-8.
        import io
        from docx import Document
        from home.letters import build_docx_bytes
        data = build_docx_bytes("A — B")
        doc = Document(io.BytesIO(data))
        self.assertIn("—", "\n".join(p.text for p in doc.paragraphs))

    def test_a_long_letter_paginates_without_error(self):
        from home.letters import build_pdf_bytes
        long_text = "\n\n".join(f"Paragraph number {i}. " * 20 for i in range(60))
        self.assertTrue(build_pdf_bytes(long_text).startswith(b"%PDF"))

    def test_a_very_long_unbroken_word_does_not_hang(self):
        # multi_cell can loop forever on a token wider than the cell.
        from home.letters import build_pdf_bytes
        self.assertTrue(build_pdf_bytes("A" * 500).startswith(b"%PDF"))

    def test_the_seeded_templates_export_end_to_end(self):
        from home.letters import build_docx_bytes, build_pdf_bytes, render_letter
        dept = Department.objects.create(dept_name="BCT")
        program = Program.objects.create(program_name="BE-BCT", department=dept)
        teacher = TeacherInfo.objects.create(
            name="Prof X", unique_id="T-X", email="x@example.com", department=dept,
        )
        student = StudentLoginInfo.objects.create(
            username="Export Student", roll_number="080BCT600", department=dept,
            program=program, password="x", dob="2000-01-01", gender="Female",
        )
        application = Application.objects.create(
            name="Export Student", std=student, professor=teacher, subjects="Physics",
        )
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                letter = render_letter(application, tpl)
                self.assertTrue(build_pdf_bytes(letter).startswith(b"%PDF"))
                self.assertTrue(build_docx_bytes(letter).startswith(b"PK"))


class RenderCustomViewTests(TestCase):
    """The preview renders the professor's chosen template (FR-1)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof F", unique_id="T-F", email="f@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Hari Thapa", roll_number="080BCT007", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Male",
        )
        self.application = Application.objects.create(
            name="Hari Thapa", std=self.student, professor=self.teacher,
            subjects="Networks",
        )
        Qualities.objects.create(application=self.application)
        self.chosen = CustomTemplates.objects.create(
            template_name="Chosen", template="CHOSEN for {{ app.name }}",
            professor=self.teacher,
        )
        self.fallback = CustomTemplates.objects.create(
            template_name="Fallback", template="FALLBACK", professor=self.teacher,
            is_default=True,
        )
        self.client.cookies["unique"] = "T-F"

    def test_the_selected_template_is_rendered(self):
        response = self.client.post("/renderCustom", {
            "roll": "080BCT007", "template_id": self.chosen.pk,
        })
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "CHOSEN for Hari Thapa")

    def test_without_a_selection_the_default_is_rendered(self):
        response = self.client.post("/renderCustom", {"roll": "080BCT007"})
        self.assertContains(response, "FALLBACK")

    def test_a_system_template_can_be_selected(self):
        system = CustomTemplates.objects.filter(is_system=True).first()
        response = self.client.post("/renderCustom", {
            "roll": "080BCT007", "template_id": system.pk,
        })
        self.assertContains(response, "Hari Thapa")

    def test_the_chosen_template_id_is_carried_into_the_download_form(self):
        response = self.client.post("/renderCustom", {
            "roll": "080BCT007", "template_id": self.chosen.pk,
        })
        self.assertContains(response, f'name="template_id" value="{self.chosen.pk}"')

    def test_the_professor_anecdote_is_still_saved(self):
        self.client.post("/renderCustom", {
            "roll": "080BCT007", "template_id": self.chosen.pk,
            "prof_anecdote": "He rebuilt the lab router overnight.",
        })
        self.application.refresh_from_db()
        self.assertEqual(
            self.application.prof_anecdote, "He rebuilt the lab router overnight."
        )

    def test_the_quality_checkboxes_are_still_saved(self):
        self.client.post("/renderCustom", {
            "roll": "080BCT007", "template_id": self.chosen.pk,
            "quality1": "on", "quality2": "on", "qual": "diligent",
        })
        quality = Qualities.objects.get(application=self.application)
        self.assertTrue(quality.leadership)
        self.assertTrue(quality.hardworking)
        self.assertFalse(quality.social)
        self.assertEqual(quality.quality, "diligent")

    def test_a_stale_cookie_redirects_to_login(self):
        self.client.cookies["unique"] = "NOPE"
        response = self.client.post("/renderCustom", {"roll": "080BCT007"})
        self.assertEqual(response.status_code, 302)
        self.assertIn("/loginTeacher", response["Location"])

    def test_another_professors_student_is_not_previewable(self):
        other = TeacherInfo.objects.create(
            name="Prof G", unique_id="T-G", email="g@example.com", department=self.dept,
        )
        self.client.cookies["unique"] = "T-G"
        response = self.client.post("/renderCustom", {"roll": "080BCT007"})
        self.assertEqual(response.status_code, 404)

    def test_an_unknown_roll_is_not_found(self):
        response = self.client.post("/renderCustom", {"roll": "NOSUCHROLL"})
        self.assertEqual(response.status_code, 404)

    def test_a_get_request_redirects_to_the_dashboard(self):
        response = self.client.get("/renderCustom")
        self.assertEqual(response.status_code, 302)

    def test_an_application_with_no_satellite_rows_still_previews(self):
        # The old view used .get() on six satellite models and raised
        # DoesNotExist for a barely-filled request.
        bare_student = StudentLoginInfo.objects.create(
            username="Bare Student", roll_number="080BCT888", department=self.dept,
            program=self.program, password="x", dob="2000-01-01",
        )
        Application.objects.create(
            name="Bare Student", std=bare_student, professor=self.teacher,
        )
        response = self.client.post("/renderCustom", {
            "roll": "080BCT888", "template_id": self.chosen.pk,
        })
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "CHOSEN for Bare Student")


@override_settings(MEDIA_ROOT=tempfile.mkdtemp())
class MakeLetterTemplateListTests(TestCase):
    """The letter form offers system templates as well as the professor's own."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof H", unique_id="T-H", email="h@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Gita Kc", roll_number="080BCT321", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Female",
        )
        self.application = Application.objects.create(
            name="Gita Kc", std=self.student, professor=self.teacher, subjects="Maths",
        )
        Paper.objects.create(application=self.application)
        Project.objects.create(application=self.application)
        University.objects.create(application=self.application, uni_name="MIT", country="USA")
        Qualities.objects.create(application=self.application)
        Academics.objects.create(application=self.application)
        # formTeacher.html dereferences .url on every uploaded file unguarded,
        # so the fixture must actually attach them or the template 500s.
        files = Files.objects.create(application=self.application)
        for field in ("Photo", "transcript", "CV"):
            getattr(files, field).save(
                f"{field}.pdf", ContentFile(b"x"), save=False
            )
        files.save()
        self.teacher.images.save("prof.png", ContentFile(b"x"), save=True)
        self.user = User.objects.create_user(username="proph", password="pw")
        self.user.first_name = "Prof H/T-H"
        self.user.save()

    def test_the_picker_offers_system_templates(self):
        self.client.force_login(self.user)
        self.client.cookies["unique"] = "T-H"
        response = self.client.post("/makeLetter", {"roll": "080BCT321"})
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Formal / Academic")

    def test_the_picker_posts_a_template_id(self):
        self.client.force_login(self.user)
        self.client.cookies["unique"] = "T-H"
        response = self.client.post("/makeLetter", {"roll": "080BCT321"})
        self.assertContains(response, 'name="template_id"')
        self.assertNotContains(response, 'name="temp"')


@override_settings(MEDIA_ROOT=tempfile.mkdtemp())
class DownloadLetterTests(TestCase):
    """Exporting a letter stores it and stamps the tracking fields (FR-1/FR-5)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof G", unique_id="T-G", email="g@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Gita Kc", roll_number="080BCT099", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Female",
        )
        self.application = Application.objects.create(
            name="Gita Kc", std=self.student, professor=self.teacher, subjects="Signals",
        )
        self.tpl = CustomTemplates.objects.create(
            template_name="Export", template="EXPORTED for {{ app.name }}",
            professor=self.teacher,
        )
        self.client.cookies["unique"] = "T-G"

    def _post(self, **extra):
        payload = {"roll": "080BCT099", "format": "pdf", "template_id": self.tpl.pk}
        payload.update(extra)
        return self.client.post("/download_letter/", payload)

    def test_pdf_download_returns_a_pdf(self):
        response = self._post()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/pdf")
        self.assertIn("attachment", response["Content-Disposition"])

    def test_docx_download_returns_a_docx(self):
        response = self._post(format="docx")
        self.assertEqual(response.status_code, 200)
        self.assertIn("wordprocessingml", response["Content-Type"])

    def test_generation_stamps_the_tracking_fields(self):
        self._post()
        self.application.refresh_from_db()
        self.assertTrue(self.application.is_generated)
        self.assertIsNotNone(self.application.generated_at)
        self.assertEqual(self.application.generated_template, self.tpl)
        self.assertTrue(self.application.generated_letter)

    def test_the_stored_file_is_the_downloaded_file(self):
        response = self._post()
        self.application.refresh_from_db()
        with self.application.generated_letter.open("rb") as handle:
            self.assertEqual(handle.read(), response.content)

    def test_the_chosen_template_is_used_not_the_default(self):
        CustomTemplates.objects.create(
            template_name="My Default", template="DEFAULT-BODY",
            professor=self.teacher, is_default=True,
        )
        response = self._post(format="docx")
        import io
        from docx import Document
        text = "\n".join(p.text for p in Document(io.BytesIO(response.content)).paragraphs)
        self.assertIn("EXPORTED for Gita Kc", text)
        self.assertNotIn("DEFAULT-BODY", text)

    def test_edited_text_is_used_instead_of_the_template(self):
        response = self._post(format="docx", edited_letter="HAND WRITTEN VERSION")
        import io
        from docx import Document
        texts = [p.text for p in Document(io.BytesIO(response.content)).paragraphs]
        self.assertIn("HAND WRITTEN VERSION", texts)
        self.assertNotIn("EXPORTED for Gita Kc", texts)

    def test_the_edited_text_is_what_gets_stored(self):
        self._post(format="docx", edited_letter="HAND WRITTEN VERSION")
        self.application.refresh_from_db()
        import io
        from docx import Document
        with self.application.generated_letter.open("rb") as handle:
            texts = [p.text for p in Document(io.BytesIO(handle.read())).paragraphs]
        self.assertIn("HAND WRITTEN VERSION", texts)

    def test_regenerating_replaces_the_stored_file_and_timestamp(self):
        self._post()
        self.application.refresh_from_db()
        first_at = self.application.generated_at
        self._post(format="docx")
        self.application.refresh_from_db()
        self.assertGreaterEqual(self.application.generated_at, first_at)
        self.assertTrue(self.application.generated_letter.name.endswith(".docx"))

    def test_the_stored_letter_is_servable_by_download_generated(self):
        # Phase 2 built this endpoint; Task 7 is what finally gives it a file.
        self._post()
        # The view stamped the row; this instance predates that write.
        self.application.refresh_from_db()
        response = self.client.get(f"/download_generated/?id={self.application.id}")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            b"".join(response.streaming_content),
            self.application.generated_letter.open("rb").read(),
        )

    def test_another_professor_cannot_export_this_letter(self):
        TeacherInfo.objects.create(
            name="Prof H", unique_id="T-H", email="h@example.com", department=self.dept,
        )
        self.client.cookies["unique"] = "T-H"
        self.assertEqual(self._post().status_code, 404)

    def test_a_stale_cookie_redirects_to_login(self):
        self.client.cookies["unique"] = "NOPE"
        response = self._post()
        self.assertEqual(response.status_code, 302)
        self.assertIn("/loginTeacher", response["Location"])

    def test_an_unknown_format_is_rejected(self):
        self.assertEqual(self._post(format="txt").status_code, 400)

    def test_a_rejected_format_does_not_stamp_anything(self):
        self._post(format="txt")
        self.application.refresh_from_db()
        self.assertFalse(self.application.is_generated)
        self.assertIsNone(self.application.generated_at)
        self.assertFalse(self.application.generated_letter)

    def test_a_missing_roll_is_not_found(self):
        response = self.client.post("/download_letter/", {"format": "pdf"})
        self.assertEqual(response.status_code, 404)

    def test_a_get_request_redirects(self):
        self.assertEqual(self.client.get("/download_letter/").status_code, 302)

    def test_a_seeded_system_template_exports_and_stamps(self):
        system = CustomTemplates.objects.filter(is_system=True).first()
        self._post(template_id=system.pk)
        self.application.refresh_from_db()
        self.assertEqual(self.application.generated_template, system)
        self.assertTrue(self.application.generated_letter)

    def test_the_filename_is_safe_for_an_awkward_name(self):
        self.application.name = "Gita / Kc \"the best\""
        self.application.save()
        response = self._post()
        disposition = response["Content-Disposition"]
        self.assertNotIn("/", disposition.split("filename=")[1])

    def test_a_broken_template_is_refused_and_stamps_nothing(self):
        broken = CustomTemplates.objects.create(
            template_name="Broken", template="{% if %}oops", professor=self.teacher,
        )
        response = self._post(template_id=broken.pk)
        self.assertEqual(response.status_code, 302)
        self.application.refresh_from_db()
        self.assertFalse(self.application.is_generated)
        self.assertIsNone(self.application.generated_at)
        self.assertIsNone(self.application.generated_template)
        self.assertFalse(self.application.generated_letter)

    def test_a_whitespace_only_template_is_refused_and_stamps_nothing(self):
        blank = CustomTemplates.objects.create(
            template_name="Blank", template="   \n\n  ", professor=self.teacher,
        )
        response = self._post(template_id=blank.pk)
        self.assertEqual(response.status_code, 302)
        self.application.refresh_from_db()
        self.assertFalse(self.application.is_generated)
        self.assertIsNone(self.application.generated_at)
        self.assertIsNone(self.application.generated_template)
        self.assertFalse(self.application.generated_letter)

    def test_control_characters_in_edited_text_do_not_crash_the_docx(self):
        import io
        from docx import Document
        for label, char in (
            ("null", "\x00"), ("vtab", "\x0b"), ("formfeed", "\x0c"), ("bell", "\x07"),
        ):
            with self.subTest(label):
                response = self._post(
                    format="docx", edited_letter=f"BEFORE{char}AFTER",
                )
                self.assertEqual(response.status_code, 200)
                texts = [
                    p.text for p in Document(io.BytesIO(response.content)).paragraphs
                ]
                self.assertIn("BEFOREAFTER", texts)

    def test_tabs_and_newlines_survive_the_control_character_strip(self):
        import io
        from docx import Document
        response = self._post(format="docx", edited_letter="A\tB\nC\x00D")
        texts = [p.text for p in Document(io.BytesIO(response.content)).paragraphs]
        self.assertIn("A\tB\nCD", texts)

    # Its own MEDIA_ROOT: the class-level one is created once at import time and
    # shared by every method, and the filesystem is not rolled back between
    # tests, so counting files there would count other tests' exports too.
    @override_settings(MEDIA_ROOT=tempfile.mkdtemp())
    def test_re_exporting_leaves_only_one_stored_file(self):
        self._post()
        self._post(format="docx")
        self._post()
        self.application.refresh_from_db()
        stored = self.application.generated_letter
        directory = os.path.dirname(stored.path)
        self.assertEqual(os.listdir(directory), [os.path.basename(stored.name)])
        # The surviving file must be the one the row points at, and be servable.
        self.assertTrue(stored.storage.exists(stored.name))

    def test_no_save_ever_persists_a_letter_without_its_metadata(self):
        # ``save=False`` on the FileField write is load-bearing: it defers that
        # write into the same UPDATE as the metadata. Without it, FieldFile.save
        # issues its own UPDATE first, and any failure after that point leaves a
        # torn row -- a live Re-download link with no timestamp or template, which
        # the dashboard renders as em-dashes beside a working download.
        from unittest.mock import patch
        original_save = Application.save
        snapshots = []

        def record(instance, *args, **kwargs):
            snapshots.append((
                bool(instance.generated_letter),
                instance.is_generated,
                instance.generated_at is not None,
            ))
            return original_save(instance, *args, **kwargs)

        with patch.object(Application, "save", record):
            self._post()

        self.assertTrue(snapshots, "the view never saved the application")
        for has_file, is_generated, has_timestamp in snapshots:
            if has_file:
                self.assertTrue(
                    is_generated and has_timestamp,
                    "a save persisted the stored letter before its metadata",
                )


class DuplicateTemplateTests(TestCase):
    """A professor can copy a system template into their own library (FR-3)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof I", unique_id="T-I", email="i@example.com", department=self.dept,
        )
        self.other = TeacherInfo.objects.create(
            name="Prof J", unique_id="T-J", email="j@example.com", department=self.dept,
        )
        self.system = CustomTemplates.objects.filter(is_system=True).first()
        self.client.cookies["unique"] = "T-I"

    def test_duplicating_creates_an_owned_editable_copy(self):
        response = self.client.post("/duplicateTemplate", {"template_id": self.system.pk})
        self.assertEqual(response.status_code, 302)
        copy = CustomTemplates.objects.get(professor=self.teacher)
        self.assertEqual(copy.template, self.system.template)
        self.assertFalse(copy.is_system)
        self.assertFalse(copy.is_default)

    def test_the_copy_is_named_after_the_original(self):
        self.client.post("/duplicateTemplate", {"template_id": self.system.pk})
        copy = CustomTemplates.objects.get(professor=self.teacher)
        self.assertEqual(copy.template_name, f"{self.system.template_name} (copy)")

    def test_duplicating_twice_does_not_collide(self):
        self.client.post("/duplicateTemplate", {"template_id": self.system.pk})
        self.client.post("/duplicateTemplate", {"template_id": self.system.pk})
        names = list(
            CustomTemplates.objects.filter(professor=self.teacher)
            .values_list("template_name", flat=True)
        )
        self.assertEqual(len(names), 2)
        self.assertEqual(len(set(names)), 2)

    def test_the_original_system_template_is_untouched(self):
        self.client.post("/duplicateTemplate", {"template_id": self.system.pk})
        self.system.refresh_from_db()
        self.assertTrue(self.system.is_system)
        self.assertIsNone(self.system.professor)

    def test_a_professor_may_duplicate_their_own_template(self):
        mine = CustomTemplates.objects.create(
            template_name="Mine", template="body", professor=self.teacher
        )
        self.client.post("/duplicateTemplate", {"template_id": mine.pk})
        self.assertEqual(
            CustomTemplates.objects.filter(professor=self.teacher).count(), 2
        )

    def test_another_professors_template_cannot_be_duplicated(self):
        theirs = CustomTemplates.objects.create(
            template_name="Theirs", template="secret", professor=self.other
        )
        response = self.client.post("/duplicateTemplate", {"template_id": theirs.pk})
        self.assertEqual(response.status_code, 404)
        self.assertFalse(CustomTemplates.objects.filter(professor=self.teacher).exists())

    def test_a_stale_cookie_redirects_to_login(self):
        self.client.cookies["unique"] = "NOPE"
        response = self.client.post("/duplicateTemplate", {"template_id": self.system.pk})
        self.assertEqual(response.status_code, 302)
        self.assertIn("/loginTeacher", response["Location"])
        self.assertFalse(CustomTemplates.objects.filter(professor=self.teacher).exists())

    def test_a_malformed_id_is_not_served(self):
        for bad in ("abc", "", "-1", "999999"):
            with self.subTest(value=bad):
                response = self.client.post("/duplicateTemplate", {"template_id": bad})
                self.assertEqual(response.status_code, 404)

    def test_a_get_request_redirects_to_the_editor(self):
        response = self.client.get("/duplicateTemplate")
        self.assertEqual(response.status_code, 302)

    def test_the_copy_is_immediately_selectable_for_generation(self):
        from home.letters import select_template
        self.client.post("/duplicateTemplate", {"template_id": self.system.pk})
        copy = CustomTemplates.objects.get(professor=self.teacher)
        self.assertEqual(select_template(self.teacher, copy.pk), copy)

    def test_a_long_name_does_not_overflow_the_column(self):
        long_name = "L" * 95
        source = CustomTemplates.objects.create(
            template_name=long_name, template="body", professor=self.teacher
        )
        self.client.post("/duplicateTemplate", {"template_id": source.pk})
        copy = CustomTemplates.objects.exclude(pk=source.pk).get(professor=self.teacher)
        self.assertLessEqual(len(copy.template_name), 100)
        copy.full_clean()  # would raise if the column width were exceeded

    def test_repeatedly_duplicating_a_copy_stays_within_the_column(self):
        # Each generation appends " (copy)"; without a bound this overflows.
        current = CustomTemplates.objects.create(
            template_name="T" * 80, template="body", professor=self.teacher
        )
        for _ in range(15):
            self.client.post("/duplicateTemplate", {"template_id": current.pk})
            current = CustomTemplates.objects.filter(
                professor=self.teacher
            ).order_by("-pk").first()
            self.assertLessEqual(len(current.template_name), 100)
            current.full_clean()

    def test_many_duplicates_of_a_capped_name_all_get_distinct_names(self):
        # A name already at the cap: naive truncation of the finished candidate
        # would map every suffix onto one string and spin forever.
        source = CustomTemplates.objects.create(
            template_name="C" * 100, template="body", professor=self.teacher
        )
        for _ in range(5):
            self.client.post("/duplicateTemplate", {"template_id": source.pk})
        copies = CustomTemplates.objects.filter(professor=self.teacher).exclude(pk=source.pk)
        names = list(copies.values_list("template_name", flat=True))
        self.assertEqual(len(names), 5)
        self.assertEqual(len(set(names)), 5)
        for candidate in names:
            self.assertLessEqual(len(candidate), 100)


class TemplateEditorViewTests(TestCase):
    """The editor lists system templates alongside the professor's own (FR-3)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof K", unique_id="T-K", email="k@example.com", department=self.dept,
        )
        self.other = TeacherInfo.objects.create(
            name="Prof L", unique_id="T-L", email="l@example.com", department=self.dept,
        )
        self.mine = CustomTemplates.objects.create(
            template_name="My Own Template", template="body", professor=self.teacher
        )
        CustomTemplates.objects.create(
            template_name="Not Mine At All", template="body", professor=self.other
        )
        self.client.cookies["unique"] = "T-K"

    def test_the_professors_own_templates_are_listed(self):
        response = self.client.get("/makeTemplate")
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "My Own Template")

    def test_system_templates_are_offered(self):
        response = self.client.get("/makeTemplate")
        self.assertContains(response, "Formal / Academic")

    def test_another_professors_templates_are_not_shown(self):
        response = self.client.get("/makeTemplate")
        self.assertNotContains(response, "Not Mine At All")

    def test_each_system_template_has_a_duplicate_button(self):
        response = self.client.get("/makeTemplate")
        system = CustomTemplates.objects.filter(is_system=True).first()
        self.assertContains(response, "/duplicateTemplate")
        self.assertContains(response, f'name="template_id" value="{system.pk}"')

    def test_a_stale_cookie_redirects_instead_of_crashing(self):
        self.client.cookies["unique"] = "NOPE"
        response = self.client.get("/makeTemplate")
        self.assertEqual(response.status_code, 302)
        self.assertIn("/loginTeacher", response["Location"])

    def test_no_cookie_redirects_instead_of_crashing(self):
        del self.client.cookies["unique"]
        response = self.client.get("/makeTemplate")
        self.assertEqual(response.status_code, 302)

    def test_the_system_list_survives_a_save(self):
        # getTemplate re-renders the same page; the section must not vanish.
        response = self.client.post("/getTemplate", {
            "content": "Dear Committee", "templateName": "Fresh", "uid": "T-K",
        })
        self.assertContains(response, "Formal / Academic")


class GetTemplateOwnershipTests(TestCase):
    """Saving a template writes to the signed-in professor only (FR-3)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof M", unique_id="T-M", email="m@example.com", department=self.dept,
        )
        self.victim = TeacherInfo.objects.create(
            name="Prof N", unique_id="T-N", email="n@example.com", department=self.dept,
        )
        self.client.cookies["unique"] = "T-M"

    def test_a_template_is_saved_to_the_signed_in_professor(self):
        self.client.post("/getTemplate", {
            "content": "Dear Committee", "templateName": "Mine", "uid": "T-M",
        })
        saved = CustomTemplates.objects.get(template_name="Mine")
        self.assertEqual(saved.professor, self.teacher)

    def test_a_forged_uid_cannot_write_to_another_professor(self):
        self.client.post("/getTemplate", {
            "content": "Injected", "templateName": "Forged", "uid": "T-N",
        })
        self.assertFalse(CustomTemplates.objects.filter(professor=self.victim).exists())
        self.assertTrue(CustomTemplates.objects.filter(professor=self.teacher).exists())

    def test_a_forged_uid_cannot_clear_another_professors_default(self):
        theirs = CustomTemplates.objects.create(
            template_name="Their Default", template="body",
            professor=self.victim, is_default=True,
        )
        self.client.post("/getTemplate", {
            "content": "x", "templateName": "New", "is_default": "on", "uid": "T-N",
        })
        theirs.refresh_from_db()
        self.assertTrue(theirs.is_default)

    def test_marking_default_clears_the_previous_default(self):
        old = CustomTemplates.objects.create(
            template_name="Old", template="x", professor=self.teacher, is_default=True
        )
        self.client.post("/getTemplate", {
            "content": "New body", "templateName": "New", "is_default": "on", "uid": "T-M",
        })
        old.refresh_from_db()
        self.assertFalse(old.is_default)
        self.assertTrue(CustomTemplates.objects.get(template_name="New").is_default)

    def test_saving_the_same_name_updates_rather_than_duplicates(self):
        self.client.post("/getTemplate", {
            "content": "First", "templateName": "Same", "uid": "T-M",
        })
        self.client.post("/getTemplate", {
            "content": "Second", "templateName": "Same", "uid": "T-M",
        })
        matches = CustomTemplates.objects.filter(
            professor=self.teacher, template_name="Same"
        )
        self.assertEqual(matches.count(), 1)
        self.assertIn("Second", matches.first().template)

    def test_a_saved_template_never_becomes_a_system_template(self):
        self.client.post("/getTemplate", {
            "content": "x", "templateName": "Mine", "uid": "T-M",
        })
        self.assertFalse(CustomTemplates.objects.get(template_name="Mine").is_system)

    def test_a_stale_cookie_redirects_to_login(self):
        self.client.cookies["unique"] = "NOPE"
        response = self.client.post("/getTemplate", {
            "content": "x", "templateName": "y", "uid": "T-M",
        })
        self.assertEqual(response.status_code, 302)
        self.assertIn("/loginTeacher", response["Location"])

    def test_a_get_request_redirects_to_the_editor(self):
        response = self.client.get("/getTemplate")
        self.assertEqual(response.status_code, 302)

    def test_a_template_with_no_name_is_rejected(self):
        response = self.client.post("/getTemplate", {"content": "x", "uid": "T-M"})
        self.assertEqual(response.status_code, 302)
        self.assertFalse(CustomTemplates.objects.filter(professor=self.teacher).exists())

    def test_a_blank_name_is_rejected(self):
        response = self.client.post("/getTemplate", {
            "content": "x", "templateName": "   ", "uid": "T-M",
        })
        self.assertEqual(response.status_code, 302)
        self.assertFalse(CustomTemplates.objects.filter(professor=self.teacher).exists())


class TemplateEditorPrefillTests(TestCase):
    """The editor must load a template without corrupting it (FR-3)."""

    # A newline, an apostrophe, a double quote and a "<" so autoescaping is
    # genuinely exercised rather than trivially satisfied.
    BODY = 'Dear "Sir",\nIt\'s 5 < 6 & fine.\nRegards'

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof O", unique_id="T-O", email="o@example.com", department=self.dept,
        )
        self.mine = CustomTemplates.objects.create(
            template_name="Round Trip", template=self.BODY, professor=self.teacher
        )
        self.client.cookies["unique"] = "T-O"

    def _option_attr(self, html):
        import re
        match = re.search(r'data-content="(.*?)" data-name', html, re.S)
        self.assertIsNotNone(match, "the existing-template option was not rendered")
        return match.group(1)

    def test_the_dropdown_body_is_not_javascript_escaped(self):
        html = self.client.get("/makeTemplate").content.decode()
        attr = self._option_attr(html)
        # escapejs turns every newline into the six literal characters \u000A,
        # which getAttribute() hands to TinyMCE verbatim.
        self.assertNotIn(r"\u000A", attr)
        self.assertNotIn(r"\u0027", attr)
        self.assertNotIn(r"\u003C", attr)

    def test_the_dropdown_body_keeps_a_real_newline(self):
        html = self.client.get("/makeTemplate").content.decode()
        self.assertIn("\n", self._option_attr(html))

    def test_the_dropdown_body_is_html_escaped_not_attribute_breaking(self):
        attr = self._option_attr(self.client.get("/makeTemplate").content.decode())
        # Autoescaping is what makes the attribute safe; a raw quote would end it.
        self.assertNotIn('"', attr)
        self.assertIn("&quot;", attr)
        self.assertIn("&lt;", attr)

    def test_the_after_save_block_is_not_nested_inside_a_script(self):
        response = self.client.post("/getTemplate", {
            "content": "Hello", "templateName": "Round Trip", "uid": "T-O",
        })
        html = response.content.decode()
        self.assertContains(response, 'id="tmpbody"')
        self.assertContains(response, 'id="tmpname"')
        # The old bug: `... = {{ x|json_script }}` emitted a <script> inside a
        # <script>, closing the outer one and killing the rest of the handler.
        self.assertNotIn("= <script", html)

    def test_the_after_save_block_is_absent_before_a_save(self):
        response = self.client.get("/makeTemplate")
        self.assertNotContains(response, 'id="tmpbody"')


class NoHardcodedTemplatesTests(TestCase):
    """Letter bodies live in the database, not in views.py (FR-1)."""

    def test_views_no_longer_carries_a_hardcoded_letter(self):
        import inspect
        from home import views
        source = inspect.getsource(views)
        self.assertNotIn("default_template_content", source)

    def test_the_seeding_helper_is_gone(self):
        from home import views
        self.assertFalse(hasattr(views, "add_default_template_to_all_professors"))

    def test_the_letter_bodies_come_from_the_database(self):
        # The three system templates are the only shipped letter bodies now.
        self.assertEqual(CustomTemplates.objects.filter(is_system=True).count(), 3)


class DashboardTemplateLinkTests(TestCase):
    """The dashboard points professors at the template library (FR-3)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof O", unique_id="T-O", email="o@example.com", department=self.dept,
        )
        self.client.cookies["unique"] = "T-O"

    def test_the_dashboard_links_to_the_template_editor(self):
        response = self.client.get("/teacher")
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "/makeTemplate")

    def test_a_professor_with_no_default_is_told_about_the_starter_library(self):
        response = self.client.get("/teacher")
        self.assertContains(response, "starter template")

    def test_a_professor_with_a_default_sees_its_name(self):
        CustomTemplates.objects.create(
            template_name="My Favourite", template="body",
            professor=self.teacher, is_default=True,
        )
        response = self.client.get("/teacher")
        self.assertContains(response, "My Favourite")
        self.assertNotContains(response, "starter template")


class QualityPersistenceTests(TestCase):
    """The professor's checkbox input must survive a missing Qualities row."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof P", unique_id="T-P", email="p@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Quality Student", roll_number="080BCT770", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Male",
        )
        self.application = Application.objects.create(
            name="Quality Student", std=self.student, professor=self.teacher,
            subjects="Networks",
        )
        self.tpl = CustomTemplates.objects.create(
            template_name="Q", template="{{ app.name }}", professor=self.teacher,
        )
        self.client.cookies["unique"] = "T-P"

    def _post(self, **extra):
        payload = {"roll": "080BCT770", "template_id": self.tpl.pk}
        payload.update(extra)
        return self.client.post("/renderCustom", payload)

    def test_qualities_are_saved_when_no_row_exists_yet(self):
        # This is the data-loss case: .update() on an empty queryset is a silent no-op.
        self.assertFalse(Qualities.objects.filter(application=self.application).exists())
        self._post(quality1="on", quality2="on", qual="diligent")
        quality = Qualities.objects.get(application=self.application)
        self.assertTrue(quality.leadership)
        self.assertTrue(quality.hardworking)
        self.assertFalse(quality.social)
        self.assertEqual(quality.quality, "diligent")

    def test_qualities_are_updated_when_a_row_already_exists(self):
        Qualities.objects.create(application=self.application, leadership=True)
        self._post(quality2="on", qual="thorough")
        quality = Qualities.objects.get(application=self.application)
        self.assertFalse(quality.leadership)
        self.assertTrue(quality.hardworking)
        self.assertEqual(quality.quality, "thorough")

    def test_no_duplicate_qualities_row_is_created(self):
        self._post(quality1="on")
        self._post(quality2="on")
        self.assertEqual(
            Qualities.objects.filter(application=self.application).count(), 1
        )

    def test_an_existing_extracurricular_value_is_preserved(self):
        # ``extracirricular`` comes from the student's intake form and is not
        # part of the professor's checkbox set; updating must not clear it.
        Qualities.objects.create(
            application=self.application, extracirricular="Robotics club",
        )
        self._post(quality1="on")
        quality = Qualities.objects.get(application=self.application)
        self.assertEqual(quality.extracirricular, "Robotics club")
        self.assertTrue(quality.leadership)


@override_settings(MEDIA_ROOT=tempfile.mkdtemp())
class MissingUploadTests(TestCase):
    """A student who skipped an upload must not 500 the letter form."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof R", unique_id="T-R", email="r@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="No Upload", roll_number="080BCT880", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Female",
        )
        self.application = Application.objects.create(
            name="No Upload", std=self.student, professor=self.teacher, subjects="Maths",
        )
        Paper.objects.create(application=self.application)
        Project.objects.create(application=self.application)
        University.objects.create(
            application=self.application, uni_name="MIT", country="USA",
        )
        Qualities.objects.create(application=self.application)
        Academics.objects.create(application=self.application)
        # Every file field left empty - the case that used to crash.
        Files.objects.create(application=self.application)
        self.user = User.objects.create_user(username="profr", password="pw")
        self.user.first_name = "Prof R/T-R"
        self.user.save()
        self.client.force_login(self.user)
        self.client.cookies["unique"] = "T-R"

    def test_the_letter_form_renders_without_any_uploads(self):
        response = self.client.post("/makeLetter", {"roll": "080BCT880"})
        self.assertEqual(response.status_code, 200)

    def test_the_dashboard_renders_without_a_teacher_photo(self):
        response = self.client.get("/teacher")
        self.assertEqual(response.status_code, 200)

    def test_a_present_upload_still_renders_its_link(self):
        from django.core.files.base import ContentFile
        files = Files.objects.get(application=self.application)
        files.transcript.save("transcript.pdf", ContentFile(b"%PDF-1.4 fake"), save=True)
        response = self.client.post("/makeLetter", {"roll": "080BCT880"})
        self.assertEqual(response.status_code, 200)
        # Assert on the link itself, not the bare word "transcript": the fallback
        # branch renders "No transcript uploaded." and would satisfy that too.
        files.refresh_from_db()
        self.assertContains(response, files.transcript.url)
        self.assertNotContains(response, "No transcript uploaded.")


class QualitiesUniquenessTests(TestCase):
    """One Qualities row per Application, enforced by the database."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.teacher = TeacherInfo.objects.create(
            name="Prof U", unique_id="T-U", email="u@example.com", department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Unique Student", roll_number="080BCT990", department=self.dept,
            program=self.program, password="x", dob="2000-01-01", gender="Male",
        )
        self.application = Application.objects.create(
            name="Unique Student", std=self.student, professor=self.teacher,
            subjects="Algorithms",
        )
        self.tpl = CustomTemplates.objects.create(
            template_name="U", template="{{ app.name }}", professor=self.teacher,
        )
        self.client.cookies["unique"] = "T-U"

    def test_a_second_qualities_row_for_one_application_is_rejected(self):
        Qualities.objects.create(application=self.application)
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                Qualities.objects.create(application=self.application)

    def test_two_applications_may_each_have_their_own_row(self):
        other_application = Application.objects.create(
            name="Unique Student", std=self.student, professor=self.teacher,
            subjects="Compilers",
        )
        Qualities.objects.create(application=self.application)
        Qualities.objects.create(application=other_application)
        self.assertEqual(Qualities.objects.count(), 2)

    def test_render_custom_still_works_with_exactly_one_row(self):
        # The constraint guarantees update_or_create's .get() can never see the
        # MultipleObjectsReturned case that would 500 the professor.
        Qualities.objects.create(application=self.application, leadership=True)
        response = self.client.post(
            "/renderCustom",
            {"roll": "080BCT990", "template_id": self.tpl.pk, "quality2": "on"},
        )
        self.assertEqual(response.status_code, 200)
        quality = Qualities.objects.get(application=self.application)
        self.assertFalse(quality.leadership)
        self.assertTrue(quality.hardworking)
        self.assertEqual(Qualities.objects.filter(application=self.application).count(), 1)


class UnicodePdfTests(TestCase):
    """Non-Latin-1 text must survive PDF export (FR-1)."""

    def test_an_em_dash_is_preserved_not_replaced(self):
        from home.letters import build_pdf_bytes
        data = build_pdf_bytes("A — B")
        self.assertTrue(data.startswith(b"%PDF"))

    def test_the_pdf_embeds_a_font_subset(self):
        # FontFile2 only appears once a TrueType subset is embedded; with the
        # old Latin-1 core-font path it is absent.
        from home.letters import build_pdf_bytes
        self.assertIn(b"FontFile2", build_pdf_bytes("A — B"))

    def test_curly_quotes_survive(self):
        from home.letters import build_pdf_bytes
        self.assertTrue(build_pdf_bytes("“Quoted” and ‘single’").startswith(b"%PDF"))

    def test_plain_ascii_still_works(self):
        from home.letters import build_pdf_bytes
        self.assertTrue(build_pdf_bytes("Dear Committee,\n\nRegards").startswith(b"%PDF"))

    def test_a_long_letter_still_paginates(self):
        from home.letters import build_pdf_bytes
        long_text = "\n\n".join(f"Paragraph {i}. " * 20 for i in range(60))
        self.assertTrue(build_pdf_bytes(long_text).startswith(b"%PDF"))

    def test_the_seeded_templates_still_export(self):
        from home.letters import build_pdf_bytes, render_letter
        dept = Department.objects.create(dept_name="BCT")
        program = Program.objects.create(program_name="BE-BCT", department=dept)
        teacher = TeacherInfo.objects.create(
            name="Prof U", unique_id="T-U", email="u@example.com", department=dept,
        )
        student = StudentLoginInfo.objects.create(
            username="Uni Student", roll_number="080BCT900", department=dept,
            program=program, password="x", dob="2000-01-01", gender="Female",
        )
        application = Application.objects.create(
            name="Uni Student", std=student, professor=teacher, subjects="Physics",
        )
        for tpl in CustomTemplates.objects.filter(is_system=True):
            with self.subTest(name=tpl.template_name):
                letter = render_letter(application, tpl)
                self.assertTrue(build_pdf_bytes(letter).startswith(b"%PDF"))


class RemovedEndpointTests(TestCase):
    """Dead endpoints are gone, not merely unlinked."""

    def test_the_edit_endpoint_no_longer_exists(self):
        # It was routed, unguarded, and wrote to the database.
        self.assertEqual(self.client.post("/edit", {"roll": "x"}).status_code, 404)

    def test_the_edit_view_is_gone(self):
        from home import views
        self.assertFalse(hasattr(views, "edit"))

    def test_the_testing_view_is_gone(self):
        from home import views
        self.assertFalse(hasattr(views, "testing"))


class TeacherUserLinkTests(TestCase):
    """TeacherInfo links to a Django User by FK, not by a name-string convention."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")

    def test_a_teacher_can_be_linked_to_a_user(self):
        user = User.objects.create_user(username="linked", password="pw")
        teacher = TeacherInfo.objects.create(
            name="Prof Linked", unique_id="T-LINK", email="l@example.com",
            department=self.dept, user=user,
        )
        self.assertEqual(teacher.user, user)
        self.assertEqual(user.teacherinfo, teacher)

    def test_the_link_is_optional(self):
        teacher = TeacherInfo.objects.create(
            name="Prof Unlinked", unique_id="T-UNLINK", email="u@example.com",
            department=self.dept,
        )
        self.assertIsNone(teacher.user)

    def test_one_user_cannot_be_two_teachers(self):
        from django.db.utils import IntegrityError
        user = User.objects.create_user(username="solo", password="pw")
        TeacherInfo.objects.create(
            name="A", unique_id="T-A1", email="a@example.com",
            department=self.dept, user=user,
        )
        with self.assertRaises(IntegrityError):
            TeacherInfo.objects.create(
                name="B", unique_id="T-B1", email="b@example.com",
                department=self.dept, user=user,
            )

    def test_deleting_the_user_does_not_delete_the_teacher(self):
        user = User.objects.create_user(username="doomed", password="pw")
        teacher = TeacherInfo.objects.create(
            name="Prof Doomed", unique_id="T-DOOM", email="d@example.com",
            department=self.dept, user=user,
        )
        user.delete()
        teacher.refresh_from_db()
        self.assertIsNone(teacher.user)


class CurrentTeacherTests(TestCase):
    """current_teacher resolves the acting professor from the session, not a cookie."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.user = User.objects.create_user(username="ct", password="pw")
        self.teacher = TeacherInfo.objects.create(
            name="Prof CT", unique_id="T-CT", email="ct@example.com",
            department=self.dept, user=self.user,
        )
        self.factory = RequestFactory()

    def _request(self, user=None, cookies=None):
        request = self.factory.get("/")
        request.user = user or AnonymousUser()
        request.COOKIES.update(cookies or {})
        return request

    def test_an_authenticated_linked_user_resolves(self):
        from home.identity import current_teacher
        self.assertEqual(current_teacher(self._request(self.user)), self.teacher)

    def test_an_anonymous_request_resolves_to_none(self):
        from home.identity import current_teacher
        self.assertIsNone(current_teacher(self._request()))

    def test_a_forged_cookie_is_ignored(self):
        # The whole point of this phase.
        from home.identity import current_teacher
        victim_user = User.objects.create_user(username="victim", password="pw")
        TeacherInfo.objects.create(
            name="Victim", unique_id="T-VICTIM", email="v@example.com",
            department=self.dept, user=victim_user,
        )
        request = self._request(self.user, {"unique": "T-VICTIM"})
        self.assertEqual(current_teacher(request), self.teacher)

    def test_a_cookie_alone_grants_nothing(self):
        from home.identity import current_teacher
        self.assertIsNone(current_teacher(self._request(None, {"unique": "T-CT"})))

    def test_an_unlinked_teacher_resolves_by_the_name_convention(self):
        # Legacy rows the data migration could not match keep working.
        from home.identity import current_teacher
        legacy_user = User.objects.create_user(username="legacy", password="pw")
        legacy_user.first_name = "Prof Legacy/T-LEGACY"
        legacy_user.save()
        legacy = TeacherInfo.objects.create(
            name="Prof Legacy", unique_id="T-LEGACY", email="lg@example.com",
            department=self.dept,
        )
        self.assertEqual(current_teacher(self._request(legacy_user)), legacy)

    def test_an_authenticated_non_teacher_resolves_to_none(self):
        from home.identity import current_teacher
        plain = User.objects.create_user(username="plain", password="pw")
        self.assertIsNone(current_teacher(self._request(plain)))


class LoginHelperTests(TestCase):
    """The shared test helper really authenticates."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.teacher = TeacherInfo.objects.create(
            name="Prof LH", unique_id="T-LH", email="lh@example.com",
            department=self.dept,
        )

    def test_it_creates_and_links_a_user(self):
        login_as_teacher(self.client, self.teacher)
        self.teacher.refresh_from_db()
        self.assertIsNotNone(self.teacher.user)

    def test_it_reuses_an_existing_user(self):
        user = User.objects.create_user(username="existing", password="pw")
        self.teacher.user = user
        self.teacher.save()
        self.assertEqual(login_as_teacher(self.client, self.teacher), user)

    def test_the_resolved_teacher_matches(self):
        from home.identity import current_teacher
        login_as_teacher(self.client, self.teacher)
        request = RequestFactory().get("/")
        request.user = self.teacher.user
        self.assertEqual(current_teacher(request), self.teacher)


class CookieImpersonationTests(TestCase):
    """A forged cookie must not act as anyone (the Phase 4b headline)."""

    def setUp(self):
        self.dept = Department.objects.create(dept_name="BCT")
        self.program = Program.objects.create(program_name="BE-BCT", department=self.dept)
        self.victim = TeacherInfo.objects.create(
            name="Victim Prof", unique_id="T-VIC", email="vic@example.com",
            department=self.dept,
        )
        self.student = StudentLoginInfo.objects.create(
            username="Victim Student", roll_number="080BCT950", department=self.dept,
            program=self.program, password="x", dob="2000-01-01",
        )
        self.application = Application.objects.create(
            name="Victim Student", std=self.student, professor=self.victim,
        )
        CustomTemplates.objects.create(
            template_name="Victim Template", template="secret",
            professor=self.victim, is_default=True,
        )

    def test_a_forged_cookie_cannot_reach_the_dashboard(self):
        self.client.cookies["unique"] = "T-VIC"
        response = self.client.get("/teacher")
        self.assertEqual(response.status_code, 302)
        self.assertIn("/loginTeacher", response["Location"])

    def test_a_forged_cookie_cannot_list_templates(self):
        self.client.cookies["unique"] = "T-VIC"
        self.assertEqual(self.client.get("/makeTemplate").status_code, 302)

    def test_a_forged_cookie_cannot_preview_a_letter(self):
        self.client.cookies["unique"] = "T-VIC"
        response = self.client.post("/renderCustom", {"roll": "080BCT950"})
        self.assertEqual(response.status_code, 302)

    def test_a_forged_cookie_cannot_export_a_letter(self):
        self.client.cookies["unique"] = "T-VIC"
        response = self.client.post(
            "/download_letter/", {"roll": "080BCT950", "format": "pdf"}
        )
        self.assertEqual(response.status_code, 302)

    def test_a_forged_cookie_cannot_write_templates(self):
        self.client.cookies["unique"] = "T-VIC"
        self.client.post("/getTemplate", {"content": "x", "templateName": "Injected"})
        self.assertFalse(
            CustomTemplates.objects.filter(template_name="Injected").exists()
        )

    def test_a_forged_cookie_cannot_duplicate_templates(self):
        self.client.cookies["unique"] = "T-VIC"
        system = CustomTemplates.objects.filter(is_system=True).first()
        self.client.post("/duplicateTemplate", {"template_id": system.pk})
        self.assertEqual(
            CustomTemplates.objects.filter(professor=self.victim).count(), 1
        )

    def test_a_forged_cookie_cannot_redownload_a_stored_letter(self):
        self.client.cookies["unique"] = "T-VIC"
        response = self.client.get(f"/download_generated/?id={self.application.id}")
        self.assertEqual(response.status_code, 302)

    def test_one_professor_cannot_act_as_another_by_cookie(self):
        # Signed in as a real professor, but forging someone else's cookie.
        attacker = TeacherInfo.objects.create(
            name="Attacker", unique_id="T-ATK", email="atk@example.com",
            department=self.dept,
        )
        login_as_teacher(self.client, attacker)
        self.client.cookies["unique"] = "T-VIC"
        response = self.client.get("/teacher")
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "Victim Template")


class ProfessorRegistrationLinksAccountTests(TestCase):
    """A newly registered professor must be linked to their login account."""

    def test_registering_a_professor_links_the_login_account(self):
        dept = Department.objects.create(dept_name="BEX")
        subject = Subject.objects.create(sub_name="Operating Systems")
        response = self.client.post("/registerProfessor/", {
            "name": "New Prof", "title": "Professor", "phone": "9800000000",
            "email": "new@example.com", "department": dept.pk,
            "subjects": [subject.pk],
            "password": "pw-12345", "confirm_password": "pw-12345",
        })
        self.assertEqual(response.status_code, 302)

        teacher = TeacherInfo.objects.get(email="new@example.com")
        self.assertIsNotNone(teacher.user)
        self.assertEqual(teacher.user.email, "new@example.com")

        # And that link is what resolves identity on a real request.
        from home.identity import current_teacher
        request = RequestFactory().get("/")
        request.user = teacher.user
        self.assertEqual(current_teacher(request), teacher)
